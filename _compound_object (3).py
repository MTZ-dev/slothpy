from os import path
from typing import Any
import h5py
import numpy as np
from cycler import cycler
from docx import Document as Doc
import matplotlib.pyplot as plt
from matplotlib.ticker import AutoMinorLocator, MultipleLocator
import matplotlib.colors
import matplotlib.cm
import matplotlib.gridspec
from matplotlib.animation import PillowWriter
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from slothpy.magnetism.g_tensor import calculate_g_tensor_and_axes_doublet
from slothpy.magnetism.magnetisation import (mth, mag_3d)
from slothpy.magnetism.susceptibility import (chitht, chit_tensorht, chit_3d)
from slothpy.general_utilities.grids_over_hemisphere import lebedev_laikov_grid
from slothpy.general_utilities.io import (get_soc_energies_cm_1, get_states_magnetic_momenta,
                                          get_states_total_angular_momneta,
                                          get_total_angular_momneta_matrix, get_magnetic_momenta_matrix)
from slothpy.magnetism.zeeman import (zeeman_splitting, get_zeeman_matrix)
from slothpy.angular_momentum.pseudo_spin_ito import (get_decomposition_in_z_total_angular_momentum_basis,
                                                      get_decomposition_in_z_magnetic_momentum_basis,
                                                      ito_real_decomp_matrix,
                                                      ito_complex_decomp_matrix,
                                                      get_soc_matrix_in_z_magnetic_momentum_basis,
                                                      get_soc_matrix_in_z_total_angular_momentum_basis,
                                                      get_zeeman_matrix_in_z_magnetic_momentum_basis,
                                                      get_zeeman_matrix_in_z_total_angular_momentum_basis,
                                                      matrix_from_ito_complex, matrix_from_ito_real)


class Compound:

    @classmethod
    def _new(cls, filepath: str, filename: str):

        filename += ".slt"

        hdf5_file = path.join(filepath, filename)

        obj = super().__new__(cls)

        obj._hdf5 = hdf5_file
        obj.get_hdf5_groups_and_attributes()

        return obj

    def __new__(cls, *args, **kwargs):
        raise TypeError(
            "The Compound object should not be instantiated directly. Use a Compound creation function instead.")

    def __repr__(self) -> str:

        representation = f"Compound from {self._hdf5} with the following groups of data:\n"

        for group, attributes in self._groups.items():
            representation += f"{group}: {attributes}\n"

        return representation

    def __str__(self) -> str:

        string = f"Compound from {self._hdf5} with the following groups of data:\n"

        for group, attributes in self._groups.items():
            string += f"{group}: {attributes}\n"

        return string

    def __setitem__(self, key, value) -> None:

        value = np.array(value)

        if isinstance(key, str):

            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_dataset = file.create_dataset(key, shape=value.shape, dtype=value.dtype)
                    new_dataset[:] = value[:]

                self.get_hdf5_groups_and_attributes()
                return

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to set dataset {key} in .slt file: {self._hdf5}: {error_type}: {error_message}')

        elif isinstance(key, tuple) and len(key) == 2 and isinstance(key[0], str) and isinstance(key[1], str):

            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(key[0])
                    new_dataset = new_group.create_dataset(key[1], shape=value.shape, dtype=value.dtype)
                    new_dataset[:] = value[:]

                self.get_hdf5_groups_and_attributes()
                return

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to set group "{key[0]}" and dataset "{key[1]}" in .slt file: {self._hdf5}: {error_type}: {error_message}')

        else:
            raise KeyError("Invalid key type. It has to be str or 2-tuple of str.")

    def __getitem__(self, key) -> Any:

        if isinstance(key, str):

            try:
                with h5py.File(self._hdf5, 'r') as file:

                    value = file[key][:]

                return value

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to get dataset {key} from .slt file: {self._hdf5}: {error_type}: {error_message}')

        elif isinstance(key, tuple) and len(key) == 2 and isinstance(key[0], str) and isinstance(key[1], str):

            try:
                with h5py.File(self._hdf5, 'r') as file:

                    value = file[key[0]][key[1]][:]

                return value

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to get group "{key[0]}" and dataset "{key[1]}" from .slt file: {self._hdf5}: {error_type}: {error_message}')

        else:
            raise KeyError("Invalid key type. It has to be str or 2-tuple of str.")

    # def __getattr__(self, __name: str) -> Any:
    #     pass

    # def __setattr__(self, __name: str, __value: Any) -> None:
    #     pass

    def get_hdf5_groups_and_attributes(self):

        def collect_groups(name, obj):
            if isinstance(obj, h5py.Group):
                groups_dict[name] = dict(obj.attrs)

        groups_dict = {}

        with h5py.File(self._hdf5, 'r') as file:
            file.visititems(collect_groups)

        self._groups = groups_dict

    def delete_group(self, group: str) -> None:

        try:
            with h5py.File(self._hdf5, 'r+') as file:
                del file[group]

        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to delete group {group} from .slt file: {self._hdf5}: {error_type}: {error_message}')

        self.get_hdf5_groups_and_attributes()

    def calculate_g_tensor_and_axes_doublet(self, group: str, doublets: np.ndarray, slt: str = None):

        try:
            g_tensor_list, magnetic_axes_list = calculate_g_tensor_and_axes_doublet(self._hdf5, group, doublets)

        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to compute g-tensors and main magnetic axes from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_g_tensors_axes')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing g-tensors of doublets and their magnetic axes calculated from group: {group}.'
                    tensors = new_group.create_dataset(f'{slt}_g_tensors',
                                                       shape=(g_tensor_list.shape[0], g_tensor_list.shape[1]),
                                                       dtype=np.float64)
                    tensors.attrs[
                        'Description'] = f'Dataset containing number of doublet and respective g-tensors from group {group}.'
                    axes = new_group.create_dataset(f'{slt}_axes', shape=(
                    magnetic_axes_list.shape[0], magnetic_axes_list.shape[1], magnetic_axes_list.shape[2]),
                                                    dtype=np.float64)
                    axes.attrs[
                        'Description'] = f'Dataset containing rotation matrices from initial coordinate system to magnetic axes of respective g-tensors from group: {group}.'
                    tensors[:, :] = g_tensor_list[:, :]
                    axes[:, :, :] = magnetic_axes_list[:, :, :]

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save g-tensors and magnetic axes to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return g_tensor_list, magnetic_axes_list

    def calculate_mth(self, group: str, states_cutoff: np.int64, fields: np.ndarray, grid: np.ndarray,
                      temperatures: np.ndarray, num_cpu: int, slt: str = None):

        fields = np.array(fields)
        temperatures = np.array(temperatures)

        if isinstance(grid, int):
            grid = lebedev_laikov_grid(grid)
        else:
            grid = np.array(grid)

        try:
            mth_array = mth(self._hdf5, group, states_cutoff, fields, grid, temperatures, num_cpu)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to compute M(T,H) from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_magnetisation')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing M(T,H) magnetisation calculated from group: {group}.'
                    mth_dataset = new_group.create_dataset(f'{slt}_mth', shape=(mth_array.shape[0], mth_array.shape[1]),
                                                           dtype=np.float64)
                    mth_dataset.attrs[
                        'Description'] = f'Dataset containing M(T,H) magnetisation (T - rows, H - columns) calculated from group: {group}.'
                    fields_dataset = new_group.create_dataset(f'{slt}_fields', shape=(fields.shape[0],),
                                                              dtype=np.float64)
                    fields_dataset.attrs[
                        'Description'] = f'Dataset containing magnetic field H values used in simulation of M(T,H) from group: {group}.'
                    temperatures_dataset = new_group.create_dataset(f'{slt}_temperatures',
                                                                    shape=(temperatures.shape[0],), dtype=np.float64)
                    temperatures_dataset.attrs[
                        'Description'] = f'Dataset containing temperature T values used in simulation of M(T,H) from group: {group}.'

                    mth_dataset[:, :] = mth_array[:, :]
                    fields_dataset[:] = fields[:]
                    temperatures_dataset[:] = temperatures[:]

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save M(T,H) to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return mth_array

    def calculate_chitht(self, group: str, fields: np.ndarray, states_cutoff: int, temperatures: np.ndarray,
                         num_cpu: int, num_of_points: int, delta_h: np.float64, exp: bool = False, T: bool = True,
                         grid: np.ndarray = None, slt: str = None) -> np.ndarray:

        fields = np.array(fields)
        temperatures = np.array(temperatures)

        if T:
            chi_name = 'chiT(H,T)'
            chi_file = 'chit'
        else:
            chi_name = 'chi(H,T)'
            chi_file = 'chi'

        if isinstance(grid, int):
            grid = lebedev_laikov_grid(grid)
        else:
            grid = np.array(grid)

        try:
            chitht_array = chitht(self._hdf5, group, fields, states_cutoff, temperatures, num_cpu, num_of_points,
                                  delta_h, exp, T, grid)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to compute {chi_name} from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_susceptibility')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing {chi_name} magnetic susceptibility calculated from group: {group}.'
                    chitht_dataset = new_group.create_dataset(f'{slt}_{chi_file}ht',
                                                              shape=(chitht_array.shape[0], chitht_array.shape[1]),
                                                              dtype=np.float64)
                    chitht_dataset.attrs[
                        'Description'] = f'Dataset containing {chi_name} magnetic susceptibility (H - rows, T - columns) calculated from group: {group}.'
                    fields_dataset = new_group.create_dataset(f'{slt}_fields', shape=(fields.shape[0],),
                                                              dtype=np.float64)
                    fields_dataset.attrs[
                        'Description'] = f'Dataset containing magnetic field H values used in simulation of {chi_name} from group: {group}.'
                    temperatures_dataset = new_group.create_dataset(f'{slt}_temperatures',
                                                                    shape=(temperatures.shape[0],), dtype=np.float64)
                    temperatures_dataset.attrs[
                        'Description'] = f'Dataset containing temperature T values used in simulation of {chi_name} from group: {group}.'

                    chitht_dataset[:, :] = chitht_array[:, :]
                    fields_dataset[:] = fields[:]
                    temperatures_dataset[:] = temperatures[:]

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save {chi_name} to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return chitht_array

    def calculate_chit_tensorht(self, group: str, fields: np.ndarray, states_cutoff: int, temperatures: np.ndarray,
                                num_cpu: int, num_of_points: int, delta_h: np.float64, T: bool = True, slt: str = None):

        fields = np.array(fields)
        temperatures = np.array(temperatures)

        try:
            chit_tensorht_array = chit_tensorht(self._hdf5, group, fields, states_cutoff, temperatures, num_cpu,
                                                num_of_points, delta_h, T)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to compute chi_tensor(H,T) from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_susceptibility_tensor')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing chiT_tensor(H,T) Van Vleck susceptibility tensor calculated from group: {group}.'
                    chit_tensorht_dataset = new_group.create_dataset(f'{slt}_chit_tensorht', shape=(
                    chit_tensorht_array.shape[0], chit_tensorht_array.shape[1], 3, 3), dtype=np.float64)
                    chit_tensorht_dataset.attrs[
                        'Description'] = f'Dataset containing chiT_tensor(H,T) Van Vleck susceptibility tensor (H, T, 3, 3) calculated from group: {group}.'
                    fields_dataset = new_group.create_dataset(f'{slt}_fields', shape=(fields.shape[0],),
                                                              dtype=np.float64)
                    fields_dataset.attrs[
                        'Description'] = f'Dataset containing magnetic field H values used in simulation of chiT_tensor(H,T) Van Vleck susceptibility tensor from group: {group}.'
                    temperatures_dataset = new_group.create_dataset(f'{slt}_temperatures',
                                                                    shape=(temperatures.shape[0],), dtype=np.float64)
                    temperatures_dataset.attrs[
                        'Description'] = f'Dataset containing temperature T values used in simulation of chiT_tensor(H,T) Van Vleck susceptibility tensor from group: {group}.'

                    chit_tensorht_dataset[:, :] = chit_tensorht_array[:, :]
                    fields_dataset[:] = fields[:]
                    temperatures_dataset[:] = temperatures[:]

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save chiT(H,T) to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return chit_tensorht_array

    def soc_energies_cm_1(self, group: str, num_of_states: int = None, slt: str = None) -> np.ndarray:

        try:
            soc_energies_array = get_soc_energies_cm_1(self._hdf5, group, num_of_states)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get SOC energies from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_soc_energies')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing SOC (Spin-Orbit Coupling) energies calculated from group: {group}.'
                    soc_energies_dataset = new_group.create_dataset(f'{slt}_soc_energies',
                                                                    shape=(soc_energies_array.shape[0],),
                                                                    dtype=np.float64)
                    soc_energies_dataset.attrs[
                        'Description'] = f'Dataset containing SOC (Spin-Orbit Coupling) energies calculated from group: {group}.'

                    soc_energies_dataset[:] = soc_energies_array[:]

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save SOC (Spin-Orbit Coupling) energies to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return soc_energies_array

    def states_magnetic_momenta(self, group: str, states: np.ndarray = None, rotation=None, slt: str = None):

        states = np.array(states)

        try:
            states, magnetic_momenta_array = get_states_magnetic_momenta(self._hdf5, group, states, rotation)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get states magnetic momenta from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_states_magnetic_momenta')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing states magnetic momenta calculated from group: {group}.'
                    magnetic_momenta_dataset = new_group.create_dataset(f'{slt}_magnetic_momenta', shape=(
                    magnetic_momenta_array.shape[0], magnetic_momenta_array.shape[1]), dtype=np.float64)
                    magnetic_momenta_dataset.attrs[
                        'Description'] = f'Dataset containing states magnetic momenta (0-x,1-y,2-z) calculated from group: {group}.'
                    states_dataset = new_group.create_dataset(f'{slt}_states', shape=(states.shape[0],), dtype=np.int64)
                    states_dataset.attrs[
                        'Description'] = f'Dataset containing indexes of states used in simulation of magnetic momenta from group: {group}.'

                    magnetic_momenta_dataset[:] = magnetic_momenta_array[:]
                    states_dataset[:] = states[:]

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save states magnetic momenta to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return magnetic_momenta_array

    def states_total_angular_momenta(self, group: str, states: np.ndarray = None, rotation=None, slt: str = None):

        states = np.array(states)

        try:
            states, total_angular_momenta_array = get_states_total_angular_momneta(self._hdf5, group, states, rotation)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get states total angular momenta from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_states_total_angular_momenta')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing states total angular momenta calculated from group: {group}.'
                    total_angular_momenta_dataset = new_group.create_dataset(f'{slt}_total_angular_momenta', shape=(
                    total_angular_momenta_array.shape[0], total_angular_momenta_array.shape[1]), dtype=np.float64)
                    total_angular_momenta_dataset.attrs[
                        'Description'] = f'Dataset containing states total angular momenta (0-x,1-y,2-z) calculated from group: {group}.'
                    states_dataset = new_group.create_dataset(f'{slt}_states', shape=(states.shape[0],), dtype=np.int64)
                    states_dataset.attrs[
                        'Description'] = f'Dataset containing indexes of states used in simulation of total angular momenta from group: {group}.'

                    total_angular_momenta_dataset[:] = total_angular_momenta_array[:]
                    states_dataset[:] = states[:]

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save states total angular momenta to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return total_angular_momenta_array

    def calculate_zeeman_splitting(self, group: str, states_cutoff: int, num_of_states: int, fields: np.ndarray,
                                   grid: np.ndarray, num_cpu: int, average: bool = False, slt: str = None):

        fields = np.array(fields)

        if isinstance(grid, int):
            grid = lebedev_laikov_grid(grid)
            average = True

        grid = np.array(grid)

        try:
            zeeman_array = zeeman_splitting(self._hdf5, group, states_cutoff, num_of_states, fields, grid, num_cpu,
                                            average)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to compute Zeeman splitting from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_zeeman_splitting')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing Zeeman splitting calculated from group: {group}.'
                    zeeman_splitting_dataset = new_group.create_dataset(f'{slt}_zeeman', shape=zeeman_array.shape,
                                                                        dtype=np.float64)
                    if average:
                        zeeman_splitting_dataset.attrs[
                            'Description'] = f'Dataset containing Zeeman splitting averaged over grid of directions with shape: (field, energy) calculated from group: {group}.'
                    else:
                        zeeman_splitting_dataset.attrs[
                            'Description'] = f'Dataset containing Zeeman splitting with shape: (orientation, field, energy) calculated from group: {group}.'
                    fields_dataset = new_group.create_dataset(f'{slt}_fields', shape=(fields.shape[0],),
                                                              dtype=np.float64)
                    fields_dataset.attrs[
                        'Description'] = f'Dataset containing magnetic field H values used in simulation of Zeeman splitting from group: {group}.'
                    if average:
                        orientations_dataset = new_group.create_dataset(f'{slt}_orientations',
                                                                        shape=(grid.shape[0], grid.shape[1]),
                                                                        dtype=np.float64)
                        orientations_dataset.attrs[
                            'Description'] = f'Dataset containing magnetic field orientation grid with weights used in simulation of averaged Zeeman splitting from group: {group}.'
                        orientations_dataset[:] = grid[:]
                    else:
                        orientations_dataset = new_group.create_dataset(f'{slt}_orientations', shape=(grid.shape[0], 3),
                                                                        dtype=np.float64)
                        orientations_dataset.attrs[
                            'Description'] = f'Dataset containing orientations of magnetic field used in simulation of Zeeman splitting from group: {group}.'
                        orientations_dataset[:] = grid[:, :3]

                    zeeman_splitting_dataset[:, :] = zeeman_array[:, :]
                    fields_dataset[:] = fields[:]

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save Zeeman splitting to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return zeeman_array

    def total_angular_momenta_matrix(self, group: str, states_cutoff: np.int64 = None, rotation=None, slt: str = None):

        if (not isinstance(states_cutoff, np.int)) or (states_cutoff < 0):
            raise ValueError(f'Invalid states cutoff, set it to positive integer or 0 for all states.')

        try:
            total_angular_momenta_matrix_array = get_total_angular_momneta_matrix(self._hdf5, group, states_cutoff,
                                                                                  rotation)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get total angular momenta matrix from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_total_angular_momenta_matrix')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing total angular momenta calculated from group: {group}.'
                    total_angular_momenta_matrix_dataset = new_group.create_dataset(
                        f'{slt}_total_angular_momenta_matrix', shape=total_angular_momenta_matrix_array.shape,
                        dtype=np.complex128)
                    total_angular_momenta_matrix_dataset.attrs[
                        'Description'] = f'Dataset containing total angular momenta matrix (0-x, 1-y, 2-z) calculated from group: {group}.'
                    states_dataset = new_group.create_dataset(f'{slt}_states',
                                                              shape=(total_angular_momenta_matrix_array.shape[1],),
                                                              dtype=np.int64)
                    states_dataset.attrs[
                        'Description'] = f'Dataset containing states indexes of total angular momenta matrix from group: {group}.'

                    total_angular_momenta_matrix_dataset[:] = total_angular_momenta_matrix_array[:]
                    states_dataset[:] = np.arange(total_angular_momenta_matrix_array.shape[1], dtype=np.int64)

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save total angular momenta matrix to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return total_angular_momenta_matrix_array

    def magnetic_momenta_matrix(self, group: str, states_cutoff: np.ndarray = None, rotation=None, slt: str = None):

        if (not isinstance(states_cutoff, np.int)) or (states_cutoff < 0):
            raise ValueError(f'Invalid states cutoff, set it to positive integer or 0 for all states.')

        try:
            magnetic_momenta_matrix_array = get_magnetic_momenta_matrix(self._hdf5, group, states_cutoff, rotation)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get total angular momenta matrix from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_magnetic_momenta_matrix')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing magnetic momenta calculated from group: {group}.'
                    magnetic_momenta_matrix_dataset = new_group.create_dataset(f'{slt}_magnetic_momenta_matrix',
                                                                               shape=magnetic_momenta_matrix_array.shape,
                                                                               dtype=np.complex128)
                    magnetic_momenta_matrix_dataset.attrs[
                        'Description'] = f'Dataset containing magnetic momenta matrix (0-x, 1-y, 2-z) calculated from group: {group}.'
                    states_dataset = new_group.create_dataset(f'{slt}_states',
                                                              shape=(magnetic_momenta_matrix_array.shape[1],),
                                                              dtype=np.int64)
                    states_dataset.attrs[
                        'Description'] = f'Dataset containing states indexes of magnetic momenta matrix from group: {group}.'

                    magnetic_momenta_matrix_dataset[:] = magnetic_momenta_matrix_array[:]
                    states_dataset[:] = np.arange(magnetic_momenta_matrix_array.shape[1], dtype=np.int64)

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save magnetic momenta matrix to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return magnetic_momenta_matrix_array

    def decomposition_in_z_magnetic_momentum_basis(self, group, start_state, stop_state, rotation=None,
                                                   slt: str = None):

        if (not isinstance(stop_state, int)) or (stop_state < 0):
            raise ValueError(f'Invalid states number, set it to positive integer or 0 for all states.')

        try:
            decomposition = get_decomposition_in_z_magnetic_momentum_basis(self._hdf5, group, start_state, stop_state,
                                                                           rotation)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get decomposition in "z" magnetic momentum basis of SOC matrix from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_magnetic_decomposition')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing decomposition in "z" magnetic momentum basis of SOC matrix calculated from group: {group}.'
                    decomposition_dataset = new_group.create_dataset(f'{slt}_magnetic_momenta_matrix',
                                                                     shape=decomposition.shape, dtype=np.float64)
                    decomposition_dataset.attrs[
                        'Description'] = f'Dataset containing % decomposition (rows - SO-states, columns - basis) in "z" magnetic momentum basis of SOC matrix from group: {group}.'
                    states_dataset = new_group.create_dataset(f'{slt}_pseudo_spin_states',
                                                              shape=(decomposition.shape[0],), dtype=np.float64)
                    states_dataset.attrs[
                        'Description'] = f'Dataset containing Sz pseudo-spin states corresponding to the decomposition of SOC matrix from group: {group}.'

                    decomposition_dataset[:] = decomposition[:]
                    dim = (decomposition.shape[1] - 1) / 2
                    states_dataset[:] = np.arange(-dim, dim + 1, step=1, dtype=np.float64)

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save decomposition in "z" magnetic momentum basis of SOC matrix to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return decomposition

    def decomposition_in_z_angular_momentum_basis(self, group, start_state, stop_state, rotation=None, slt: str = None):

        if (not isinstance(stop_state, int)) or (stop_state < 0):
            raise ValueError(f'Invalid states number, set it to positive integer or 0 for all states.')

        try:
            decomposition = get_decomposition_in_z_total_angular_momentum_basis(self._hdf5, group, start_state,
                                                                                stop_state, rotation)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get decomposition in "z" total angular momentum basis of SOC matrix from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_total_angular_decomposition')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing decomposition in "z" total angular momentum basis of SOC matrix calculated from group: {group}.'
                    decomposition_dataset = new_group.create_dataset(f'{slt}_magnetic_momenta_matrix',
                                                                     shape=decomposition.shape, dtype=np.float64)
                    decomposition_dataset.attrs[
                        'Description'] = f'Dataset containing % decomposition (rows SO-states, columns - basis) in "z" total angular momentum basis of SOC matrix from group: {group}.'
                    states_dataset = new_group.create_dataset(f'{slt}_pseudo_spin_states',
                                                              shape=(decomposition.shape[0],), dtype=np.float64)
                    states_dataset.attrs[
                        'Description'] = f'Dataset containing Sz pseudo-spin states corresponding to the decomposition of SOC matrix from group: {group}.'

                    decomposition_dataset[:] = decomposition[:]
                    dim = (decomposition.shape[1] - 1) / 2
                    states_dataset[:] = np.arange(-dim, dim + 1, step=1, dtype=np.float64)

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save decomposition in "z" total angular momentum basis of SOC matrix to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return decomposition

    def soc_crystal_field_parameters(self, group, start_state, stop_state, order, even_order: bool = True,
                                     imaginary: bool = False, magnetic: bool = False, rotation=None, slt: str = None):

        if magnetic:
            try:
                soc_matrix = get_soc_matrix_in_z_magnetic_momentum_basis(self._hdf5, group, start_state, stop_state,
                                                                         rotation)
            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to get SOC matrix in "z" magnetic momentum basis from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        else:
            try:
                soc_matrix = get_soc_matrix_in_z_total_angular_momentum_basis(self._hdf5, group, start_state,
                                                                              stop_state, rotation)
            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to get SOC matrix in "z" total angular momentum basis from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        dim = (soc_matrix.shape[1] - 1) / 2

        if order > 2 * dim:
            raise ValueError(f'Order of ITO parameters exeeds 2S. Set it less or equal.')

        if imaginary:
            try:
                cfp = ito_complex_decomp_matrix(soc_matrix, order, even_order)
            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to ITO decompose SOC matrix in "z" magnetic momentum basis from file: {self._hdf5} - group {group}: {error_type}: {error_message}')
        else:
            try:
                cfp = ito_real_decomp_matrix(soc_matrix, order, even_order)
            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to ITO decompose SOC matrix in "z" total angular momentum basis from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        cfp_return = cfp

        if slt is not None:

            cfp = np.array(cfp)

            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_soc_ito_decomposition')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing ITO decomposition in "z" pseudo-spin basis of SOC matrix calculated from group: {group}.'
                    cfp_dataset = new_group.create_dataset(f'{slt}_ito_parameters', shape=cfp.shape, dtype=cfp.dtype)
                    cfp_dataset.attrs[
                        'Description'] = f'Dataset containing ITO decomposition in "z" pseudo-spin basis of SOC matrix from group: {group}.'
                    states_dataset = new_group.create_dataset(f'{slt}_pseudo_spin_states', shape=(1,), dtype=np.float64)
                    states_dataset.attrs[
                        'Description'] = f'Dataset containing S pseudo-spin number corresponding to the decomposition of SOC matrix from group: {group}.'

                    cfp_dataset[:] = cfp[:]
                    states_dataset[:] = dim

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save ITO decomposition in "z" pseudo-spin basis of SOC matrix to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return cfp_return

    def zeeman_matrix_ito_decpomosition(self, group, start_state, stop_state, field, orientation, order,
                                        imaginary: bool = False, magnetic: bool = False, rotation=None,
                                        slt: str = None):

        if magnetic:
            try:
                zeeman_matrix = get_zeeman_matrix_in_z_magnetic_momentum_basis(self._hdf5, group, field, orientation,
                                                                               start_state, stop_state, rotation)
            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to get Zeeman matrix in "z" magnetic momentum basis from file: {self._hdf5} - group {group}: {error_type}: {error_message}')
        else:
            try:
                zeeman_matrix = get_zeeman_matrix_in_z_total_angular_momentum_basis(self._hdf5, group, field,
                                                                                    orientation, start_state,
                                                                                    stop_state, rotation)
            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to get Zeeman matrix in "z" total angular momentum basis from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        dim = (zeeman_matrix.shape[1] - 1) / 2

        if order > 2 * dim:
            raise ValueError(f'Order of ITO parameters exeeds 2S. Set it less or equal.')

        if imaginary:
            try:
                cfp = ito_complex_decomp_matrix(zeeman_matrix, order)
            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to ITO decompose Zeeman matrix in "z" magnetic momentum basis from file: {self._hdf5} - group {group}: {error_type}: {error_message}')
        else:
            try:
                cfp = ito_real_decomp_matrix(zeeman_matrix, order)
            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to ITO decompose Zeeman matrix in "z" total angular momentum basis from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        cfp_return = cfp

        if slt is not None:

            cfp = np.array(cfp)

            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_zeeman_ito_decomposition')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing ITO decomposition in "z" pseudo-spin basis of Zeeman matrix calculated from group: {group}.'
                    cfp_dataset = new_group.create_dataset(f'{slt}_ito_parameters', shape=cfp.shape, dtype=cfp.dtype)
                    cfp_dataset.attrs[
                        'Description'] = f'Dataset containing ITO decomposition in "z" pseudo-spin basis of Zeeman matrix from group: {group}.'
                    states_dataset = new_group.create_dataset(f'{slt}_pseudo_spin_states', shape=(1,), dtype=np.float64)
                    states_dataset.attrs[
                        'Description'] = f'Dataset containing S pseudo-spin number corresponding to the decomposition of Zeeman matrix from group: {group}.'

                    cfp_dataset[:] = cfp[:]
                    states_dataset[:] = dim

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save ITO decomposition in "z" pseudo-spin basis of Zeeman matrix to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return cfp_return

    def zeeman_matrix(self, group: str, states_cutoff, field, orientation, slt: str = None):

        if (not isinstance(states_cutoff, np.int)) or (states_cutoff < 0):
            raise ValueError(f'Invalid states cutoff, set it to positive integer or 0 for all states.')

        try:
            zeeman_matrix_array = get_zeeman_matrix(self._hdf5, group, states_cutoff, field, orientation)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get Zeeman matrix from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_zeeman_matrix')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing Zeeman matrix calculated from group: {group}.'
                    zeeman_matrix_dataset = new_group.create_dataset(f'{slt}_zeeman_matrix',
                                                                     shape=zeeman_matrix_array.shape,
                                                                     dtype=np.complex128)
                    zeeman_matrix_dataset.attrs[
                        'Description'] = f'Dataset containing Zeeman matrix calculated from group: {group}.'
                    states_dataset = new_group.create_dataset(f'{slt}_states', shape=(zeeman_matrix_array.shape[1],),
                                                              dtype=np.int64)
                    states_dataset.attrs[
                        'Description'] = f'Dataset containing states indexes of Zeeman matrix from group: {group}.'

                    zeeman_matrix_dataset[:] = zeeman_matrix_array[:]
                    states_dataset[:] = np.arange(zeeman_matrix_array.shape[1], dtype=np.int64)

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save Zeeman matrix to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return zeeman_matrix_array

    def matrix_from_ito(self, name, imaginary: bool = False, dataset: str = None, pseudo_spin: str = None,
                        slt: str = None, matrix_type: str = None):

        if (dataset is not None) and (pseudo_spin is not None) and pseudo_spin > 0:

            try:
                J = pseudo_spin
                coefficients = self[f'{name}', f'{dataset}']
                if imaginary:
                    matrix = matrix_from_ito_complex(J, coefficients)
                else:
                    matrix = matrix_from_ito_real(J, coefficients)

            except Exception as e:
                error_type_1 = type(e).__name__
                error_message_1 = str(e)
                error_print_1 = f"{error_type_1}: {error_message_1}"
                raise Exception(
                    f'Failed to form matrix from ITO parameters.\n Error(s) encountered while trying compute the matrix: {error_print_1}')

        else:

            try:
                J = self[f'{name}_zeeman_ito_decomposition', f'{name}_pseudo_spin_states']
                coefficients = self[f'{name}_zeeman_ito_decomposition', f'{name}_ito_parameters']

            except Exception as e:
                error_type_2 = type(e).__name__
                error_message_2 = str(e)
                error_print_2 = f"{error_type_2}: {error_message_2}"
                try:
                    J = self[f'{name}_soc_ito_decomposition', f'{name}_pseudo_spin_states']
                    coefficients = self[f'{name}_soc_ito_decomposition', f'{name}_ito_parameters']

                except Exception as e:
                    error_type_3 = type(e).__name__
                    error_message_3 = str(e)
                    error_print_3 = f"{error_type_3}: {error_message_3}"
                    raise Exception(
                        f'Failed to form matrix from ITO parameters.\n Error(s) encountered while trying compute the matrix: {error_print_2}, {error_print_3}')

                else:
                    J_result = J[0]
                    if imaginary:
                        matrix = matrix_from_ito_complex(J[0], coefficients)
                    else:
                        matrix = matrix_from_ito_real(J[0], coefficients)

            else:
                J_result = J[0]
                if imaginary:
                    matrix = matrix_from_ito_complex(J[0], coefficients)
                else:
                    matrix = matrix_from_ito_real(J[0], coefficients)

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_matrix')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing matrix from ITO calculated from group: {name}.'
                    matrix_dataset = new_group.create_dataset(f'{slt}_matrix', shape=matrix.shape, dtype=np.complex128)
                    matrix_dataset.attrs[
                        'Description'] = f'Dataset containing matrix from ITO calculated from group: {name}.'
                    states_dataset = new_group.create_dataset(f'{slt}_pseudo_spin_states', shape=(1,), dtype=np.float64)
                    states_dataset.attrs[
                        'Description'] = f'Dataset containing S pseudo-spin number corresponding to the matrix from group: {name}.'

                    matrix_dataset[:] = matrix[:]
                    states_dataset[:] = J_result

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save matrix from ITO to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return matrix

    def soc_zeem_in_angular_magnetic_momentum_basis(self, group, start_state, stop_state, matrix_type, basis_type,
                                                    rotation=None, field=None, orientation=None, slt: str = None):

        if (matrix_type not in ['zeeman', 'soc']) or (basis_type not in ['angular', 'magnetic']):
            raise ValueError(f'Only valid matrix_type are "soc" or "zeeman" and basis_type are "angular" or "magnetic"')

        if matrix_type == 'zeeman' and ((field is None) or (orientation is None)):
            raise ValueError(f'For Zeeman matrix provide filed value and orientation.')

        try:
            if matrix_type == 'zeeman':
                if basis_type == 'angular':
                    matrix = get_zeeman_matrix_in_z_total_angular_momentum_basis(self._hdf5, group, field, orientation,
                                                                                 start_state, stop_state, rotation)
                elif basis_type == 'magnetic':
                    matrix = get_zeeman_matrix_in_z_magnetic_momentum_basis(self._hdf5, group, field, orientation,
                                                                            start_state, stop_state, rotation)
            elif matrix_type == 'soc':
                if basis_type == 'angular':
                    matrix = get_soc_matrix_in_z_total_angular_momentum_basis(self._hdf5, group, start_state,
                                                                              stop_state, rotation)
                elif basis_type == 'magnetic':
                    matrix = get_soc_matrix_in_z_magnetic_momentum_basis(self._hdf5, group, start_state, stop_state,
                                                                         rotation)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get {matrix_type} matrix from file in {basis_type} momentum basis: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_{matrix_type}_matrix_in_{basis_type}_basis')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing {matrix_type} matrix in {basis_type} momentum "z" basis calculated from group: {group}.'
                    matrix_dataset = new_group.create_dataset(f'{slt}_matrix', shape=matrix.shape, dtype=np.complex128)
                    matrix_dataset.attrs[
                        'Description'] = f'Dataset containing {matrix_type} matrix in {basis_type} momentum "z" basis calculated from group: {group}.'
                    states_dataset = new_group.create_dataset(f'{slt}_states', shape=(matrix.shape[1],), dtype=np.int64)
                    states_dataset.attrs[
                        'Description'] = f'Dataset containing states indexes of {matrix_type} matrix in {basis_type} momentum "z" basis from group: {group}.'

                    matrix_dataset[:] = matrix[:]
                    states_dataset[:] = np.arange(matrix.shape[1], dtype=np.int64)

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save {matrix_type} matrix in {basis_type} momentum "z" basis to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return matrix

    def calculate_mag_3d(self, group: str, states_cutoff: int, field: np.ndarray, spherical_grid: int,
                         temperature: np.float64, num_cpu: int, slt: str = None):

        temperature = np.array([temperature], dtype=np.float64)

        try:
            x, y, z = mag_3d(self._hdf5, group, states_cutoff, field, spherical_grid, temperature, num_cpu)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to compute 3D magnetisation from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:
            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_3d_magnetisation')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing 3D magnetisation calculated from group: {group}.'
                    mag_3d_dataset = new_group.create_dataset(f'{slt}_mag_3d', shape=(3, x.shape[1], x.shape[2]),
                                                              dtype=np.float64)
                    mag_3d_dataset.attrs[
                        'Description'] = f'Dataset containing 3D magnetisation as meshgird (0-x,1-y,2-z) arrays over sphere (T: {temperature} K, H: {field} T) calculated from group: {group}.'

                    mag_3d_dataset[0, :, :] = x[0]
                    mag_3d_dataset[1, :, :] = y[0]
                    mag_3d_dataset[2, :, :] = z[0]

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save 3D magnetisation to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return x, y, z

    def calculate_chit_3d(self, group: str, field: np.float64, states_cutoff: int, temperature: np.float64,
                          num_cpu: int, num_of_points: int, delta_h: np.float64, spherical_grid: int, exp: bool = False,
                          T: bool = True, slt: str = None):

        temperature = np.array([temperature], dtype=np.float64)

        try:
            x, y, z = chit_3d(self._hdf5, group, field, states_cutoff, temperature, num_cpu, num_of_points, delta_h,
                              spherical_grid, exp, T)
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to compute 3D magnetic susceptibility from file: {self._hdf5} - group {group}: {error_type}: {error_message}')

        if slt is not None:

            if T:
                chi_file = 'chit'
            else:
                chi_file = 'chi'

            try:
                with h5py.File(self._hdf5, 'r+') as file:
                    new_group = file.create_group(f'{slt}_3d_susceptibility')
                    new_group.attrs[
                        'Description'] = f'Group({slt}) containing 3D magnetic susceptibility calculated from group: {group}.'
                    chit_3d_dataset = new_group.create_dataset(f'{slt}_{chi_file}_3d',
                                                               shape=(3, x.shape[1], x.shape[2]), dtype=np.float64)
                    chit_3d_dataset.attrs[
                        'Description'] = f'Dataset containing 3D magnetic susceptibility as meshgird (0-x,1-y,2-z) arrays over sphere (T: {temperature} K, H: {field} T) calculated from group: {group}.'

                    chit_3d_dataset[0, :, :] = x[0]
                    chit_3d_dataset[1, :, :] = y[0]
                    chit_3d_dataset[2, :, :] = z[0]

                self.get_hdf5_groups_and_attributes()

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to save 3D magneic susceptibility to file: {self._hdf5} - group {slt}: {error_type}: {error_message}')

        return x, y, z

    @staticmethod
    def colour_map(name):
        """
        Creates matplotlib colour map object.

        Args:
            name: (str or lst) one of defined names for colour maps: BuPi, rainbow, dark_rainbow, light_rainbow,
            light_rainbow_alt, BuOr, BuYl, BuRd, GnYl, PrOr, GnRd, funmat, NdCoN322bpdo, NdCoNO222bpdo, NdCoI22bpdo,
            viridis, plasma, inferno, magma, cividis or list of colour from which colour map will be created by
            interpolation of colours between ones on a list; for predefined names modifiers can be applyed: _l loops
            the list in a way that it starts and ends with the same colour, _r reverses the list

        Returns:

        """
        cmap_list = []
        reverse = False
        loop = False
        if name[-2:] == '_l':
            name = name[:-2]
            loop = True
        try:
            if name[-2:] == '_r':
                reverse = True
                name = name[:-2]
            if type(name) == list:
                cmap_list = name
            elif name == 'BuPi':
                cmap_list = ['#0091ad', '#1780a1', '#2e6f95', '#455e89', '#5c4d7d', '#723c70', '#a01a58',
                             '#b7094c']
            elif name == 'rainbow':
                cmap_list = ['#ff0000', '#ff8700', '#ffd300', '#deff0a', '#a1ff0a', '#0aff99', '#0aefff', '#147df5',
                             '#580aff', '#be0aff']
            elif name == 'dark_rainbow':
                cmap_list = ['#F94144', '#F3722C', '#F8961E', '#F9844A', '#F9C74F', '#90BE6D', '#43AA8B', '#4D908E',
                             '#577590', '#277DA1']
            elif name == 'light_rainbow':
                cmap_list = ['#FFADAD', '#FFD6A5', '#FDFFB6', '#CAFFBF', '#9BF6FF', '#A0C4FF', '#BDB2FF', '#FFC6FF']
            elif name == 'light_rainbow_alt':
                cmap_list = ['#FBF8CC', '#FDE4CF', '#FFCFD2', '#F1C0E8', '#CFBAF0', '#A3C4F3', '#90DBF4', '#8EECF5',
                             '#98F5E1', '#B9FBC0']
            elif name == 'BuOr':
                cmap_list = ['#03045e', '#023e8a', '#0077b6', '#0096c7',
                             '#00b4d8', '#ff9e00',
                             '#ff9100', '#ff8500', '#ff6d00',
                             '#ff5400']
            elif name == 'BuRd':
                cmap_list = ['#033270', '#1368aa', '#4091c9', '#9dcee2',
                             '#fedfd4', '#f29479', '#ef3c2d', '#cb1b16',
                             '#65010c']
            elif name == 'BuYl':
                cmap_list = ['#184e77', '#1e6091', '#1a759f', '#168aad',
                             '#34a0a4',
                             '#52b69a', '#76c893', '#99d98c', '#b5e48c',
                             '#d9ed92']
            elif name == 'GnYl':
                cmap_list = ['#007f5f', '#2b9348', '#55a630', '#80b918',
                             '#aacc00', '#bfd200', '#d4d700', '#dddf00',
                             '#eeef20', '#ffff3f']
            elif name == 'PrOr':
                cmap_list = ['#240046', '#3c096c', '#5a189a', '#7b2cbf',
                             '#9d4edd', '#ff9e00', '#ff9100', '#ff8500',
                             '#ff7900', '#ff6d00']
            elif name == 'GnRd':
                cmap_list = ['#005C00', '#2D661B', '#2A850E', '#27A300',
                             '#A9FFA5', '#FFA5A5', '#FF0000', '#BA0C0C',
                             '#751717', '#5C0000']
            elif name == 'funmat':
                cmap_list = ['#1f6284', '#277ba5', '#2f94c6', '#49a6d4',
                             '#6ab6dc', '#ffe570', '#ffe15c', '#ffda33',
                             '#ffd20a', '#e0b700']
            elif name == 'NdCoN322bpdo':
                cmap_list = ['#00268f', '#0046ff', '#009cf4', '#E5E4E2', '#ede76d',
                             '#ffb900', '#b88700']
            elif name == 'NdCoNO222bpdo':
                cmap_list = ['#A90F97', '#E114C9', '#f9bbf2', '#77f285',
                             '#11BB25', '#0C831A']
            elif name == 'NdCoI22bpdo':
                cmap_list = ['#075F5F', '#0B9898', '#0fd1d1', '#FAB3B3',
                             '#d10f0f', '#720808']
            if cmap_list:
                if reverse:
                    cmap_list.reverse()
                if loop:
                    new_cmap_list = cmap_list.copy()
                    for i in range(len(cmap_list)):
                        new_cmap_list.append(cmap_list[-(i + 1)])
                    cmap_list = new_cmap_list
                cmap = matplotlib.colors.LinearSegmentedColormap.from_list("", cmap_list)
            elif name == 'viridis':
                cmap = matplotlib.cm.viridis
                if reverse:
                    cmap = matplotlib.cm.viridis_r
            elif name == 'plasma':
                cmap = matplotlib.cm.plasma
                if reverse:
                    cmap = matplotlib.cm.plasma_r
            elif name == 'inferno':
                cmap = matplotlib.cm.inferno
                if reverse:
                    cmap = matplotlib.cm.inferno_r
            elif name == 'magma':
                cmap = matplotlib.cm.magma
                if reverse:
                    cmap = matplotlib.cm.magma_r
            elif name == 'cividis':
                cmap = matplotlib.cm.cividis
                if reverse:
                    cmap = matplotlib.cm.cividis_r
            else:
                print(f'''There is no such colour map as {name} use one of those: BuPi, rainbow, dark_rainbow, light_rainbow, 
            light_rainbow_alt, BuOr, BuYl, BuRd, GnYl, PrOr, GnRd, funmat, NdCoN322bpdo, NdCoNO222bpdo, NdCoI22bpdo,
            viridis, plasma, inferno, magma, cividis or enter list of colours''')
            return cmap
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to find palette/colour map: {error_type}: {error_message}')

    @staticmethod
    def custom_colour_cycler(number_of_colours: int, cmap1: str, cmap2: str):
        """
        Creates colour cycler from two colour maps in alternating pattern, suitable for use in matplotlib plots.

        Args:
            number_of_colours: (int) number of colour in cycle
            cmap1: (str or lst) colour map name or list of colours (valid input for Compoud.colour_map())
            cmap2: (str or lst) colour map name or list of colours (valid input for Compoud.colour_map())

        Returns:
            cycler object created based on two input colourmaps
        """
        if number_of_colours % 2 == 0:
            increment = 0
            lst1 = Compound.colour_map(cmap1)(np.linspace(0, 1, int(number_of_colours / 2)))
            lst2 = Compound.colour_map(cmap2)(np.linspace(0, 1, int(number_of_colours / 2)))
            colour_cycler_list = []
            while increment < number_of_colours:
                if increment % 2 == 0:
                    colour_cycler_list.append(lst1[int(increment / 2)])
                else:
                    colour_cycler_list.append(lst2[int((increment - 1) / 2)])
                increment += 1
        else:
            increment = 0
            lst1 = Compound.colour_map(cmap1)(np.linspace(0, 1, int((number_of_colours / 2) + 1)))
            lst2 = Compound.colour_map(cmap2)(np.linspace(0, 1, int(number_of_colours / 2)))
            colour_cycler_list = []
            while increment < number_of_colours:
                if increment % 2 == 0:
                    colour_cycler_list.append(lst1[int(increment / 2)])
                else:
                    colour_cycler_list.append(lst2[int((increment - 1) / 2)])
                increment += 1
        return cycler(color=colour_cycler_list)

    def plot_mth(self, group: str, show=True, origin=False, save=False, colour_map_name='rainbow', xlim=(), ylim=(),
                 xticks=1, yticks=0, field='B'):
        """
        Function that creates graphs of M(H,T) given name of the group in HDF5 file, graphs can be optionally shown,
        saved, colour palettes can be changed. If origin=True it returns data packed into a dictionary for exporting
        to Origin.

            Args:
                group (str): name of a group in HDF5 file
                show (bool): determines if matplotlib graph is created
                and shown if True
                origin (bool): determines if function should return raw data
                save (bool): determines if matplotlib graph should be saved, saved graphs are TIFF files
                colour_map_name (str) or (list): sets colours used to create graphs, valid options are returned by
                Compound.colour_map staticmethod
                xlim (tuple): tuple of two or one numbers that set corresponding axe limits
                ylim (tuple): tuple of two or one numbers that set corresponding axe limits
                xticks (int): frequency of x major ticks
                yticks (int): frequency of x major ticks
                field ('B' or 'H'): chooses field type and unit: Tesla for B and kOe for H
            Returns:
                if origin=True:
                    dict[origin_column (str), data (np.array)]: contains data used to create graph in origin
        """
        try:
            """Getting data from hdf5 or sloth file"""
            mth = self[f'{group}_magnetisation', f'{group}_mth']
            fields = self[f'{group}_magnetisation', f'{group}_fields']
            if field == 'H':
                fields *= 10
                xticks *= 10
            temps = self[f'{group}_magnetisation', f'{group}_temperatures']
            """Creates dataset suitable to be exported to Origin"""
            data = {'data_x': fields, 'data_y': mth, 'comment': temps}
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get data to create graph of M(H, T): {self._hdf5} - group {group}: {error_type}: {error_message}')
        if show:
            try:
                """Plotting in matplotlib"""
                fig, ax = plt.subplots()
                """Defining colour maps for graphs"""
                colour = iter(Compound.colour_map(colour_map_name)(np.linspace(0, 1, len(temps))))
                """Creating a plot"""
                for i, mh in enumerate(mth):
                    c = next(colour)
                    ax.plot(fields, mh, linewidth=2, c=c, label=f'{temps[i]} K')

                if yticks:
                    ax.yaxis.set_major_locator(MultipleLocator(yticks))
                ax.xaxis.set_major_locator(MultipleLocator(xticks))
                ax.xaxis.set_minor_locator(AutoMinorLocator(2))
                ax.yaxis.set_minor_locator(AutoMinorLocator(2))
                ax.tick_params(which='major', length=7)
                ax.tick_params(which='minor', length=3.5)
                if field == 'B':
                    ax.set_xlabel(r'$B\ /\ \mathrm{T}$')
                elif field == 'H':
                    ax.set_xlabel(r'$H\ /\ \mathrm{kOe}$')
                ax.set_ylabel(r'$M\ /\ \mathrm{\mu_{B}}$')
                if xlim:
                    if len(xlim) == 2:
                        ax.set_ylim(xlim[0], xlim[1])
                    else:
                        ax.set_ylim(xlim[0])
                else:
                    if len(temps) > 17:
                        ax.set_xlim(0)
                        ax.legend(loc='center left', bbox_to_anchor=(1, 0.5))

                    else:
                        ax.set_xlim(0, fields[-1] + 0.3 * fields[-1])
                        ax.legend()
                if ylim:
                    if len(ylim) == 2:
                        ax.set_ylim(ylim[0], ylim[1])
                    else:
                        ax.set_ylim(ylim[0])
                else:
                    ax.set_ylim(0)
                plt.tight_layout()
                plt.show()
            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to create graph of M(H, T): {self._hdf5} - group {group}: {error_type}: {error_message}')
            if save:
                try:
                    """Saving plot figure"""
                    fig.savefig(f'mgh_{group}.tiff', dpi=300)
                except Exception as e:
                    error_type = type(e).__name__
                    error_message = str(e)
                    raise Exception(
                        f'Error encountered while trying to save graph of M(H, T): {self._hdf5} - group {group}: {error_type}: {error_message}')
        if origin:
            return data

    def plot_chitht(self, group, show=True, origin=False, save=False, colour_map_name='funmat', xlim=(), ylim=(),
                    xticks=100, yticks=0, field='B'):
        """
        Function that creates graphs of chiT(H,T) or chi(H,T) depending on content of HDF5 file, given name of the group
        in HDF5 file, graphs can be optionally shown, saved, colour palettes can be changed. If origin=True it returns
        data packed into a dictionary for exporting to Origin.

            Args:
                group (str): name of a group in HDF5 file
                show (bool): determines if matplotlib graph is created
                and shown if True
                origin (bool): determines if function should return raw data
                save (bool): determines if matplotlib graph should be saved, saved graphs are TIFF files
                colour_map_name (str): sets colours used to create graphs, valid options are returned by
                Compound.colour_map staticmethod
                xlim (tuple): tuple of two or one numbers that set corresponding axe limits
                ylim (tuple): tuple of two or one numbers that set corresponding axe limits
                xticks (int): frequency of x major ticks
                yticks (int): frequency of x major ticks
                field ('B' or 'H'): chooses field type and unit: Tesla for B and kOe for H
            Returns:
                if origin=True:
                    dict[origin_column (str), data (np.array)]: contains data used to create graph in origin
        """
        try:
            """Getting data from hdf5 or sloth file"""
            try:
                chi = self[f'{group}_susceptibility', f'{group}_chiht']
                T = False
            except:
                chi = self[f'{group}_susceptibility', f'{group}_chitht']
                T = True
            fields = self[f'{group}_susceptibility', f'{group}_fields']
            if field == 'H':
                fields *= 10
            temps = self[f'{group}_susceptibility', f'{group}_temperatures']
            """Creates dataset suitable to be exported to Origin"""
            data = {'data_x': temps, 'data_y': chi, 'comment': fields, 'T': T}
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get data to create graph of chiT(H,T) or chi(H,T): {self._hdf5} - group {group}: {error_type}: {error_message}')
        if show:
            try:
                """Plotting in matplotlib"""
                fig, ax = plt.subplots()
                """Defining colour maps for graphs"""
                colour = iter(Compound.colour_map(colour_map_name)(np.linspace(0, 1, len(fields))))
                """Creating a plot"""
                for i, ch in enumerate(chi):
                    c = next(colour)
                    ax.plot(temps, ch, linewidth=2, c=c,
                            label=f'{round(fields[i], 2)} {"kOe" if field == "H" else "T"}')
                ax.xaxis.set_major_locator(MultipleLocator(xticks))
                if yticks:
                    ax.yaxis.set_major_locator(MultipleLocator(yticks))
                ax.xaxis.set_minor_locator(AutoMinorLocator(2))
                ax.yaxis.set_minor_locator(AutoMinorLocator(2))
                ax.tick_params(which='major', length=7)
                ax.tick_params(which='minor', length=3.5)
                ax.set_xlabel(r'$T\ /\ \mathrm{K}$')
                if T:
                    ax.set_ylabel(r'$\chi_{\mathrm{M}}T\ /\ \mathrm{cm^{3}mol^{-1}K}$')
                else:
                    ax.set_ylabel(r'$\chi_{\mathrm{M}}\ /\ \mathrm{cm^{3}mol^{-1}}$')
                if xlim:
                    if len(xlim) == 2:
                        ax.set_ylim(xlim[0], xlim[1])
                    else:
                        ax.set_ylim(xlim[0])
                else:
                    ax.set_xlim(0, temps[-1])
                if ylim:
                    if len(ylim) == 2:
                        ax.set_ylim(ylim[0], ylim[1])
                    else:
                        ax.set_ylim(ylim[0])
                else:
                    ax.set_ylim(0)
                ax.legend()
                plt.tight_layout()
                plt.show()
            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to create graph of chiT(H,T) or chi(H,T): {self._hdf5} - group {group}: {error_type}: {error_message}')
            if save:
                try:
                    """Saving plot figure"""
                    fig.savefig(f'chitht_{group}.tiff', dpi=300)
                except Exception as e:
                    error_type = type(e).__name__
                    error_message = str(e)
                    raise Exception(
                        f'Error encountered while trying to save graph of chiT(H,T) or chi(H,T): {self._hdf5} - group {group}: {error_type}: {error_message}')
        if origin:
            return data

    def plot_zeeman(self, group: str, show=True, origin=False, save=False, colour_map_name1='BuPi',
                    colour_map_name2='BuPi_r', single=False, xlim=(), ylim=(), xticks=1, yticks=0, field='B'):
        """
        Function that creates graphs of E(H,orientation) given name of the group in HDF5 file, graphs can be optionally shown,
        saved, colour palettes can be changed. If origin=True it returns data packed into a dictionary for exporting
        to Origin.

            Args:
                group (str): name of a group in HDF5 file
                show (bool): determines if matplotlib graph is created
                and shown if True
                origin (bool): determines if function should return raw data
                save (bool): determines if matplotlib graph should be saved, saved graphs are TIFF files
                colour_map_name1 (str) or (list): sets colours used to create graphs, valid options are returned by
                Compound.colour_map staticmethod
                colour_map_name2 (str) or (list): sets colours used to create graphs, valid options are returned by
                Compound.colour_map staticmethod
                single (bool): determines if graph should be created for each orientation given separately
                xlim (tuple): tuple of two or one numbers that set corresponding axe limits
                ylim (tuple): tuple of two or one numbers that set corresponding axe limits
                xticks (int): frequency of x major ticks
                yticks (int): frequency of y major ticks
                field ('B' or 'H'): chooses field type and unit: Tesla for B and kOe for H
            Returns:
                if origin=True:
                    dict[origin_column (str), data (np.array)]: contains data used to create graph in origin
        """
        try:
            """Getting data from hdf5 or sloth file"""
            zeeman = self[f'{group}_zeeman_splitting', f'{group}_zeeman']
            fields = self[f'{group}_zeeman_splitting', f'{group}_fields']
            if field == 'H':
                fields *= 10
                xticks *= 10
            orientations = self[f'{group}_zeeman_splitting', f'{group}_orientations']
            """Creates dataset suitable to be exported to Origin"""
            for i, orientation in enumerate(orientations):
                data = {f'data_x{i}': fields, 'data_y': mth, f'comment{i}': orientation}
        except Exception as e:
            error_type = type(e).__name__
            error_message = str(e)
            raise Exception(
                f'Error encountered while trying to get data to create graph of E(H,orientation): {self._hdf5} - group {group}: {error_type}: {error_message}')
        if show:
            try:
                """Plotting in matplotlib"""
                if not single:
                    number_of_plots = len(orientations)
                    if number_of_plots % 5 == 0:
                        fig = plt.figure(figsize=(16, 3.2 * (number_of_plots / 5)))
                        gs = matplotlib.gridspec.GridSpec(int(number_of_plots / 5), 5)
                        devisor = 5
                    elif number_of_plots % 3 == 0:
                        fig = plt.figure(figsize=(9.6, 3.2 * (number_of_plots / 3)))
                        gs = matplotlib.gridspec.GridSpec(int(number_of_plots / 3), 3)
                        devisor = 3
                    elif number_of_plots % 2 == 0:
                        fig = plt.figure(figsize=(6.4, 3.2 * (number_of_plots / 2)))
                        gs = matplotlib.gridspec.GridSpec(int(number_of_plots / 2), 2)
                        devisor = 2
                    else:
                        fig = plt.figure(figsize=(6.4, 3.2 * number_of_plots))
                        gs = matplotlib.gridspec.GridSpec(1, number_of_plots)
                        devisor = 1
                    """Creating a plot"""
                    for i, zee in enumerate(zeeman):
                        if i % devisor != 0:
                            plt.rc('axes', prop_cycle=Compound.custom_colour_cycler(len(zeeman[0][0]), colour_map_name1,
                                                                                    colour_map_name2))
                            multiple_plots = fig.add_subplot(gs[i // devisor, i % devisor])
                            plt.plot(fields, zee, linewidth=0.75)
                            multiple_plots.xaxis.set_major_locator(MultipleLocator(xticks * 2))
                            if yticks:
                                multiple_plots.yaxis.set_major_locator(MultipleLocator(yticks))
                            multiple_plots.xaxis.set_minor_locator(AutoMinorLocator(2))
                            multiple_plots.yaxis.set_minor_locator(AutoMinorLocator(2))
                            multiple_plots.tick_params(which='major', left=False, labelleft=False, length=7)
                            multiple_plots.tick_params(which='minor', left=False, length=3.5)
                            plt.title(f'Orientation {orientations[i]}')
                            if xlim:
                                if len(xlim) == 2:
                                    multiple_plots.set_xlim(xlim[0], xlim[1])
                                else:
                                    multiple_plots.set_xlim(xlim[0])
                            if ylim:
                                if len(ylim) == 2:
                                    multiple_plots.set_ylim(ylim[0], ylim[1])
                                else:
                                    multiple_plots.set_ylim(ylim[0])

                        else:
                            if (i // devisor) == 0:
                                plt.rc('axes',
                                       prop_cycle=Compound.custom_colour_cycler(len(zeeman[0][0]), colour_map_name1,
                                                                                colour_map_name2))
                                multiple_plots = fig.add_subplot(gs[i // devisor, i % devisor])
                                plt.plot(fields, zee, linewidth=0.75)
                                multiple_plots.xaxis.set_major_locator(MultipleLocator(xticks * 2))
                                if yticks:
                                    multiple_plots.yaxis.set_major_locator(MultipleLocator(yticks))
                                multiple_plots.xaxis.set_minor_locator(AutoMinorLocator(2))
                                multiple_plots.tick_params(which='major', length=7)
                                multiple_plots.tick_params(which='minor', length=3.5)
                                multiple_plots.yaxis.set_minor_locator(AutoMinorLocator(2))
                                plt.title(f'Orientation {orientations[i]}')
                                if xlim:
                                    if len(xlim) == 2:
                                        multiple_plots.set_xlim(xlim[0], xlim[1])
                                    else:
                                        multiple_plots.set_xlim(xlim[0])
                                if ylim:
                                    if len(ylim) == 2:
                                        multiple_plots.set_ylim(ylim[0], ylim[1])
                                    else:
                                        multiple_plots.set_ylim(ylim[0])
                            else:
                                plt.rc('axes',
                                       prop_cycle=Compound.custom_colour_cycler(len(zeeman[0][0]), colour_map_name1,
                                                                                colour_map_name2))
                                multiple_plots = fig.add_subplot(gs[i // devisor, i % devisor])
                                plt.plot(fields, zee, linewidth=0.75)
                                multiple_plots.xaxis.set_major_locator(MultipleLocator(xticks * 2))
                                if yticks:
                                    multiple_plots.yaxis.set_major_locator(MultipleLocator(yticks))
                                multiple_plots.xaxis.set_minor_locator(AutoMinorLocator(2))
                                multiple_plots.tick_params(which='major', length=7)
                                multiple_plots.tick_params(which='minor', length=3.5)
                                multiple_plots.yaxis.set_minor_locator(AutoMinorLocator(2))
                                plt.title(f'Orientation {orientations[i]}')
                                if xlim:
                                    if len(xlim) == 2:
                                        multiple_plots.set_xlim(xlim[0], xlim[1])
                                    else:
                                        multiple_plots.set_xlim(xlim[0])
                                if ylim:
                                    if len(ylim) == 2:
                                        multiple_plots.set_ylim(ylim[0], ylim[1])
                                    else:
                                        multiple_plots.set_ylim(ylim[0])
                    if field == 'B':
                        fig.supxlabel(r'$B\ /\ \mathrm{T}$')
                    if field == 'H':
                        fig.supxlabel(r'$H\ /\ \mathrm{kOe}$')
                    fig.supylabel(r'$\mathrm{Energy\ /\ cm^{-1}}$')
                    plt.tight_layout()
                    plt.show()
                elif single:
                    for i, zee in enumerate(zeeman):
                        plt.rc('axes', prop_cycle=Compound.custom_colour_cycler(len(zeeman[0][0]), colour_map_name1,
                                                                                colour_map_name2))
                        fig, ax = plt.subplots()
                        ax.plot(fields, zee, linewidth=0.75)
                        plt.title(f'Orientation {orientations[i]}')
                        if field == 'B':
                            ax.set_xlabel(r'$B\ /\ \mathrm{T}$')
                        elif field == 'H':
                            ax.set_xlabel(r'$H\ /\ \mathrm{kOe}$')
                        ax.set_ylabel(r'$\mathrm{Energy\ /\ cm^{-1}}$')
                        ax.tick_params(which='major', length=7)
                        ax.tick_params(which='minor', length=3.5)
                        ax.xaxis.set_major_locator(MultipleLocator(xticks))
                        if yticks:
                            ax.yaxis.set_major_locator(MultipleLocator(yticks))
                        ax.xaxis.set_minor_locator(AutoMinorLocator(2))
                        ax.yaxis.set_minor_locator(AutoMinorLocator(2))
                        if xlim:
                            if len(xlim) == 2:
                                ax.set_xlim(xlim[0], xlim[1])
                            else:
                                ax.set_xlim(xlim[0])
                        if ylim:
                            if len(ylim) == 2:
                                ax.set_ylim(ylim[0], ylim[1])
                            else:
                                ax.set_ylim(ylim[0])
                        plt.tight_layout()
                        plt.show()
                        if save:
                            try:
                                """Saving plot figure"""
                                fig.savefig(f'zeeman_{group}_Orientation {orientations[i]}.tiff', dpi=300)
                            except Exception as e:
                                error_type = type(e).__name__
                                error_message = str(e)
                                raise Exception(
                                    f'Error encountered while trying to save graph of E(H,orientation): {self._hdf5} - group {group}: {error_type}: {error_message}')

            except Exception as e:
                error_type = type(e).__name__
                error_message = str(e)
                raise Exception(
                    f'Error encountered while trying to create graph of E(H,orientation): {self._hdf5} - group {group}: {error_type}: {error_message}')
            if save and not single:
                try:
                    """Saving plot figure"""
                    fig.savefig(f'zeeman_{group}.tiff', dpi=300)
                except Exception as e:
                    error_type = type(e).__name__
                    error_message = str(e)
                    raise Exception(
                        f'Error encountered while trying to save graph of E(H,orientation): {self._hdf5} - group {group}: {error_type}: {error_message}')
        if origin:
            return data

    def plot_chit_3d(self, group: str, show=True, save=False, colour_map_name='dark_rainbow_r', lim_scalar=1, ticks=1,
                     r_density=0, c_density=0, axis_off=False):
        """
        Creates plot of chi(t)(x,y,z)

        Args:
            group: (str) name of parent group in hdf file containing information about chi(t) dependance on direction
            show: (bool) if True shows the plot
            save: (bool) if True saves the plot
            colour_map_name: (str or list) sets colours used to create graphs, valid options are returned by
            Compound.colour_map staticmethod
            lim_scalar: (float) number by which limit of x, y and z axes will be scaled, 1 is default
            ticks: (int) spacing of major ticks on x, y and z axes
            r_density: (int) density of vertical lines building 3d image
            c_density: (int) density of horizontal lines building 3d image
            axis_off: (bool) if True axes are not shown

        Returns:

        """
        try:
            x = self[f'{group}_3d_susceptibility', f'{group}_chit_3d'][0, :, :]
            y = self[f'{group}_3d_susceptibility', f'{group}_chit_3d'][1, :, :]
            z = self[f'{group}_3d_susceptibility', f'{group}_chit_3d'][2, :, :]
            with h5py.File(self._hdf5, 'r') as file:
                description = file[f'{group}_3d_susceptibility'][f'{group}_chit_3d'].attrs['Description']
            T = True
        except:
            x = self[f'{group}_3d_susceptibility', f'{group}_chi_3d'][0, :, :]
            y = self[f'{group}_3d_susceptibility', f'{group}_chi_3d'][1, :, :]
            z = self[f'{group}_3d_susceptibility', f'{group}_chi_3d'][2, :, :]
            with h5py.File(self._hdf5, 'r') as file:
                description = file[f'{group}_3d_susceptibility'][f'{group}_chi_3d'].attrs['Description']
            T = False
        title = 'S' + description[description.index('sphere') + 1:description.index('calculated') - 1]
        if show:
            fig = plt.figure()
            ax = fig.add_subplot(projection='3d')
            max_array = np.array([np.max(x), np.max(y), np.max(z)])
            lim = np.max(max_array)
            norm = plt.Normalize(z.min(), z.max())
            colors = Compound.colour_map(colour_map_name)(norm(z))
            rcount, ccount, _ = colors.shape
            if not r_density:
                r_density = rcount
            if not c_density:
                c_density = ccount
            surface = ax.plot_surface(x, y, z, rcount=r_density, ccount=c_density, facecolors=colors, shade=False, )
            surface.set_facecolor((0, 0, 0, 0))
            ax.set_xlim(-lim * lim_scalar, lim * lim_scalar)
            ax.set_ylim(-lim * lim_scalar, lim * lim_scalar)
            ax.set_zlim(-lim * lim_scalar, lim * lim_scalar)
            # Important order of operations!
            if T:
                ax.set_xlabel(r'$\chi_{\mathrm{M}}T\ /\ \mathrm{cm^{3}mol^{-1}K}$', labelpad=20 * len(str(ticks)) / 4)
                ax.set_ylabel(r'$\chi_{\mathrm{M}}T\ /\ \mathrm{cm^{3}mol^{-1}K}$', labelpad=20 * len(str(ticks)) / 4)
                ax.set_zlabel(r'$\chi_{\mathrm{M}}T\ /\ \mathrm{cm^{3}mol^{-1}K}$', labelpad=20 * len(str(ticks)) / 4)
                # ax.set(xlabel=r'$\chi_{\mathrm{M}}T\ /\ \mathrm{cm^{3}mol^{-1}K}$', ylabel=r'$\chi_{\mathrm{M}}T\ /\ \mathrm{cm^{3}mol^{-1}K}$',zlabel=r'$\chi_{\mathrm{M}}T\ /\ \mathrm{cm^{3}mol^{-1}K}$')
            else:
                ax.set_xlabel(r'$\chi_{\mathrm{M}}\ /\ \mathrm{cm^{3}mol^{-1}}$', labelpad=20 * len(str(ticks)) / 4)
                ax.set_ylabel(r'$\chi_{\mathrm{M}}\ /\ \mathrm{cm^{3}mol^{-1}}$', labelpad=20 * len(str(ticks)) / 4)
                ax.set_zlabel(r'$\chi_{\mathrm{M}}\ /\ \mathrm{cm^{3}mol^{-1}}$', labelpad=20 * len(str(ticks)) / 4)
            if ticks == 0:
                for axis in [ax.xaxis, ax.yaxis, ax.zaxis]:
                    axis.set_ticklabels([])
                    axis._axinfo['axisline']['linewidth'] = 1
                    axis._axinfo['axisline']['color'] = (0, 0, 0)
                    axis._axinfo['grid']['linewidth'] = 0.5
                    axis._axinfo['grid']['linestyle'] = "-"
                    axis._axinfo['grid']['color'] = (0, 0, 0)
                    axis._axinfo['tick']['inward_factor'] = 0.0
                    axis._axinfo['tick']['outward_factor'] = 0.0
                    axis.set_pane_color((0.95, 0.95, 0.95))
            else:
                ax.xaxis.set_minor_locator(AutoMinorLocator(2))
                ax.yaxis.set_minor_locator(AutoMinorLocator(2))
                ax.zaxis.set_minor_locator(AutoMinorLocator(2))
                if not (not T and ticks == 1):
                    ax.xaxis.set_major_locator(MultipleLocator(ticks))
                    ax.yaxis.set_major_locator(MultipleLocator(ticks))
                    ax.zaxis.set_major_locator(MultipleLocator(ticks))
            ax.grid(False)

            ax.set_box_aspect([1, 1, 1])
            plt.title(title)
            if axis_off:
                plt.axis('off')
            plt.tight_layout()
            plt.show()

            if save:
                if axis_off:
                    fig.savefig(f'{group}_3d_chit.tiff', transparent=True, dpi=600)
                fig.savefig(f'{group}_3d_chit.tiff', dpi=600)

    def plot_mag_3d(self, group: str, show=True, save=False, colour_map_name='light_rainbow_l', lim_scalar=1,
                    ticks=0, r_density=0, c_density=0, axis_off=False):
        """
        Creates plot of magnetization(x,y,z)

        Args:
            group: (str) name of parent group in hdf file containing information about magnetisation dependance on direction
            show: (bool) if True shows the plot
            save: (bool) if True saves the plot
            colour_map_name: (str or list) sets colours used to create graphs, valid options are returned by
            Compound.colour_map staticmethod
            lim_scalar: (float) number by which limit of x, y and z axes will be scaled, 1 is default
            ticks: (int) spacing of major ticks on x, y and z axes
            r_density: (int) density of vertical lines building 3d image
            c_density: (int) density of horizontal lines building 3d image
            axis_off: (bool) if True axes are not shown

        Returns:

        """
        x = self[f'{group}_3d_magnetisation', f'{group}_mag_3d'][0, :, :]
        y = self[f'{group}_3d_magnetisation', f'{group}_mag_3d'][1, :, :]
        z = self[f'{group}_3d_magnetisation', f'{group}_mag_3d'][2, :, :]
        with h5py.File(self._hdf5, 'r') as file:
            description = file[f'{group}_3d_magnetisation'][f'{group}_mag_3d'].attrs['Description']
        title = 'S' + description[description.index('sphere') + 1:description.index('calculated') - 1]
        if show:
            fig = plt.figure()
            ax = fig.add_subplot(projection='3d')
            max_array = np.array([np.max(x), np.max(y), np.max(z)])
            lim = np.max(max_array)
            norm = plt.Normalize(z.min(), z.max())
            colors = Compound.colour_map(colour_map_name)(norm(z))
            rcount, ccount, _ = colors.shape
            if not r_density:
                r_density = rcount
            if not c_density:
                c_density = ccount
            surface = ax.plot_surface(x, y, z, rcount=r_density, ccount=c_density, facecolors=colors, shade=False, )
            surface.set_facecolor((0, 0, 0, 0))
            ax.set_xlim(-lim * lim_scalar, lim * lim_scalar)
            ax.set_ylim(-lim * lim_scalar, lim * lim_scalar)
            ax.set_zlim(-lim * lim_scalar, lim * lim_scalar)
            if ticks == 0:
                for axis in [ax.xaxis, ax.yaxis, ax.zaxis]:
                    axis.set_ticklabels([])
                    axis._axinfo['axisline']['linewidth'] = 1
                    axis._axinfo['axisline']['color'] = (0, 0, 0)
                    axis._axinfo['grid']['linewidth'] = 0.5
                    axis._axinfo['grid']['linestyle'] = "-"
                    axis._axinfo['grid']['color'] = (0, 0, 0)
                    axis._axinfo['tick']['inward_factor'] = 0.0
                    axis._axinfo['tick']['outward_factor'] = 0.0
                    axis.set_pane_color((0.95, 0.95, 0.95))
            else:
                ax.xaxis.set_minor_locator(AutoMinorLocator(2))
                ax.yaxis.set_minor_locator(AutoMinorLocator(2))
                ax.zaxis.set_minor_locator(AutoMinorLocator(2))
                ax.xaxis.set_major_locator(MultipleLocator(ticks))
                ax.yaxis.set_major_locator(MultipleLocator(ticks))
                ax.zaxis.set_major_locator(MultipleLocator(ticks))
            ax.grid(False)
            ax.set_xlabel(r'$M\ /\ \mathrm{\mu_{B}}$', labelpad=10 * len(str(ticks)) / 4)
            ax.set_ylabel(r'$M\ /\ \mathrm{\mu_{B}}$', labelpad=10 * len(str(ticks)) / 4)
            ax.set_zlabel(r'$M\ /\ \mathrm{\mu_{B}}$', labelpad=10 * len(str(ticks)) / 4)
            ax.set_box_aspect([1, 1, 1])
            if axis_off:
                plt.axis('off')
            plt.title(title)
            plt.show()

            if save:
                if axis_off:
                    fig.savefig(f'{group}_3d_mag.tiff', transparent=True, dpi=600)
                else:
                    fig.savefig(f'{group}_3d_mag.tiff', dpi=600)

    def table_energy_and_g(self, group, last_level, first_level=0, krames=True, decomp='angular', first_composition=0,
                           last_composition=0, threshold=0.5, output_name=''):
        """
        Creates .docx file containing table with g tensor and energy information

        Args:
            group: (str) name of parent group in h5py file that contains information about energy and g tensors
            last_level: (int) number of last energy level taken for creation of table counting from ground state
            first_level: (int) number of first energy level taken for creation of table counting from ground state
            krames: (bool) determines if table should use table's template for Krames lanthanide (even atomic number) or
            for non-Krames
            decomp: (str: 'angular' or 'magnetic') determines type of decomposition used
            first_composition: (int) determines first doublet which decomposition should be shown
            last_composition: (int) determines last doublet which decomposition should be shown
            threshold: (float) determines % until which components of doublet won't be shown
            output_name: (str) determines name of file

        Returns:

        """
        if not first_composition:
            first_composition = first_level
        if not last_composition:
            last_composition = last_level
        if not output_name:
            output_name = f'{group}_table_energy_and_g'
        energies = self[f'{group}_soc_energies', f'{group}_soc_energies']
        g_tensors = self[f'{group}_g_tensors_axes', f'{group}_g_tensors']
        if decomp == 'angular':
            composition_frac_matrix = self[f'{group}_total_angular_decomposition', f'{group}_magnetic_momenta_matrix']
            composition_states = self[f'{group}_total_angular_decomposition', f'{group}_pseudo_spin_states']
        elif decomp == 'magnetic':
            composition_frac_matrix = self[f'{group}_magnetic_decomposition', f'{group}_magnetic_momenta_matrix']
            composition_states = self[f'{group}_magnetic_decomposition', f'{group}_pseudo_spin_states']

        number_of_doublets = last_level - first_level + 1
        number_to_decompose = last_composition - first_composition + 1
        docx = Doc()
        if krames:
            table = docx.add_table(rows=4 + number_of_doublets, cols=5, style='Light Shading')
            row = 0

            table.cell(row, 0).paragraphs[0].add_run(f'{group}').bold = True
            table.cell(row, 0).merge(table.cell(row, 4))
            row += 1

            table.cell(row, 0).paragraphs[0].add_run(f'Energy and pseudo-')
            table.cell(row, 0).paragraphs[0].add_run(f'g').italic = True
            table.cell(row, 0).paragraphs[0].add_run('-tensor components (')
            table.cell(row, 0).paragraphs[0].add_run(f'g').italic = True
            subscript = table.cell(row, 0).paragraphs[0].add_run('x')
            subscript.font.subscript = True
            table.cell(row, 0).paragraphs[0].add_run(', ')
            table.cell(row, 0).paragraphs[0].add_run(f'g').italic = True
            subscript = table.cell(row, 0).paragraphs[0].add_run('y')
            subscript.font.subscript = True
            table.cell(row, 0).paragraphs[0].add_run(', ')
            table.cell(row, 0).paragraphs[0].add_run(f'g').italic = True
            subscript = table.cell(row, 0).paragraphs[0].add_run('z')
            subscript.font.subscript = True
            table.cell(row, 0).paragraphs[0].add_run(
                f') of {number_of_doublets}{" ground" if first_level == 0 else ""} Kramers doublets')
            table.cell(row, 0).merge(table.cell(row, 4))
            row += 1

            table.cell(row, 0).paragraphs[0].add_run(f'Doublet no.')

            table.cell(row, 1).paragraphs[0].add_run(f'Energy / cm')
            superscript = table.cell(row, 1).paragraphs[0].add_run('-1')
            superscript.font.superscript = True

            table.cell(row, 2).paragraphs[0].add_run(f'Pseudo-')
            table.cell(row, 2).paragraphs[0].add_run('g').italic = True
            table.cell(row, 2).paragraphs[0].add_run('-tensor components')
            table.cell(row, 2).merge(table.cell(row, 4))
            row += 1

            table.cell(row - 1, 0).merge(table.cell(row, 0))
            table.cell(row - 1, 1).merge(table.cell(row, 1))

            table.cell(row, 2).paragraphs[0].add_run('g').italic = True
            subscript = table.cell(row, 2).paragraphs[0].add_run('x')
            subscript.font.subscript = True

            table.cell(row, 3).paragraphs[0].add_run('g').italic = True
            subscript = table.cell(row, 3).paragraphs[0].add_run('y')
            subscript.font.subscript = True

            table.cell(row, 4).paragraphs[0].add_run('g').italic = True
            subscript = table.cell(row, 4).paragraphs[0].add_run('z')
            subscript.font.subscript = True
            row += 1

            for index in range(number_of_doublets):
                table.cell(row, 0).paragraphs[0].add_run(f'{int(g_tensors[first_level + index][0]) + 1}.')
                table.cell(row, 1).paragraphs[0].add_run(f'{energies[(first_level + index) * 2]:.3f}')
                table.cell(row, 2).paragraphs[0].add_run(f'{g_tensors[first_level + index][1]:.4f}')
                table.cell(row, 3).paragraphs[0].add_run(f'{g_tensors[first_level + index][2]:.4f}')
                table.cell(row, 4).paragraphs[0].add_run(f'{g_tensors[first_level + index][3]:.4f}')
                row += 1

            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row1 in table.rows:
                for cell in row1.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        elif not krames:
            table = docx.add_table(rows=3 + number_of_doublets, cols=4, style='Light Shading')
            row = 0

            table.cell(row, 0).paragraphs[0].add_run(f'{group}').bold = True
            table.cell(row, 0).merge(table.cell(row, 3))
            row += 1

            table.cell(row, 0).paragraphs[0].add_run(f'Energy and pseudo-')
            table.cell(row, 0).paragraphs[0].add_run(f'g').italic = True
            table.cell(row, 0).paragraphs[0].add_run('-tensor')
            table.cell(row, 0).paragraphs[0].add_run(f'g').italic = True
            subscript = table.cell(row, 0).paragraphs[0].add_run('z')
            subscript.font.subscript = True
            table.cell(row, 0).paragraphs[0].add_run(f' component')
            table.cell(row, 0).paragraphs[0].add_run(
                f') of {number_of_doublets}{" ground" if first_level == 0 else ""} Ising doublets')
            table.cell(row, 0).merge(table.cell(row, 3))
            row += 1

            table.cell(row, 0).paragraphs[0].add_run(f'Doublet no.')

            table.cell(row, 1).paragraphs[0].add_run(f'Energies / cm')
            superscript = table.cell(row, 1).paragraphs[0].add_run('-1')
            superscript.font.superscript = True

            table.cell(row, 2).paragraphs[0].add_run(f'Tunneling splitting / cm')
            superscript = table.cell(row, 2).paragraphs[0].add_run('-1')
            superscript.font.superscript = True

            table.cell(row, 3).paragraphs[0].add_run(f'g').italic = True
            subscript = table.cell(row, 3).paragraphs[0].add_run('z')
            subscript.font.subscript = True
            table.cell(row, 3).paragraphs[0].add_run(f' component')
            row += 1

            for index in range(number_of_doublets):
                table.cell(row, 0).paragraphs[0].add_run(f'{int(g_tensors[first_level + index][0]) + 1}.')
                table.cell(row, 1).paragraphs[0].add_run(f'{energies[(first_level + index * 2)]:.3f}; ')
                table.cell(row, 1).paragraphs[0].add_run(f'{energies[(first_level + 1 + index * 2)]:.3f}')
                table.cell(row, 2).paragraphs[0].add_run(
                    f'{energies[(first_level + 1 + index * 2)] - energies[(first_level + index * 2)]:.3f}')
                table.cell(row, 3).paragraphs[0].add_run(f'{g_tensors[first_level + index][2]:.4f}')
                row += 1

            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row1 in table.rows:
                for cell in row1.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        table2 = docx.add_table(rows=3, cols=number_to_decompose, style='Light Shading')
        table2.cell(0, 0).paragraphs[0].add_run(
            f'Composition of the two ground {"Kramers" if krames else "Ising"} doublets in the |')
        table2.cell(0, 0).paragraphs[0].add_run('m').italic = True
        subscript = table2.cell(0, 0).paragraphs[0].add_run('J')
        subscript.font.subscript = True
        table2.cell(0, 0).paragraphs[0].add_run('〉 basis on the quantization axes within')
        table2.cell(0, 0).paragraphs[0].add_run(' J ').italic = True
        if krames:
            table2.cell(0, 0).paragraphs[0].add_run(f'= {str(abs(composition_states[0] * 2))}/2 manifold')
        else:
            table2.cell(0, 0).paragraphs[0].add_run(f'= {str(abs(composition_states[0]))} manifold')
        table2.cell(0, 0).paragraphs[0].add_run(f' (contribution over {threshold}% shown)')
        table2.cell(0, 0).merge(table2.cell(0, number_to_decompose - 1))

        for index in range(number_to_decompose):
            table2.cell(1, index).paragraphs[0].add_run(f'{index + 1}')
            if index + 1 == 1:
                superscript = table2.cell(1, index).paragraphs[0].add_run('st')
                superscript.font.superscript = True
            elif index + 1 == 2:
                superscript = table2.cell(1, index).paragraphs[0].add_run('nd')
                superscript.font.superscript = True
            elif index + 1 == 3:
                superscript = table2.cell(1, index).paragraphs[0].add_run('rd')
                superscript.font.superscript = True
            else:
                superscript = table2.cell(1, index).paragraphs[0].add_run('th')
                superscript.font.superscript = True
            table2.cell(1, index).paragraphs[0].add_run(' doublet')
            for inner_index in range(len(composition_states)):
                if composition_frac_matrix[first_composition + index * 2][inner_index] <= threshold:
                    continue
                else:
                    table2.cell(2, index).paragraphs[0].add_run(
                        f'{composition_frac_matrix[(first_composition + index * 2)][inner_index]:.1f}%')
                    table2.cell(2, index).paragraphs[0].add_run(
                        f' |{"+" + str(int(abs(composition_states[inner_index] * 2))) if composition_states[inner_index] > 0 else "–" + str(int(abs(composition_states[inner_index] * 2)))}/2〉\n')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row2 in table2.rows:
            for cell in row2.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        docx.save(f'{output_name}.docx')

    def animate_mag_3d(self, group: str, states_cutoff: int, field: np.ndarray, spherical_grid: int,
                       temperature_start: np.float64, temperature_stop: np.float64, frames: int, num_cpu: int,
                       colour_map_name='dark_rainbow', lim_scalar=1, ticks=1, axis_off=False, r_density=0, c_density=0,
                       filename='mag_3d', fps=15, dpi=200):
        """
        Creates animation of Magnetization(x,y,z) with changing temperature
        Args:
            group: (str) name of parent group in h5py file that contains 3d magnetization information
            states_cutoff: from calculate_mag_3d
            field: from calculate_mag_3d
            spherical_grid: from calculate_mag_3d
            temperature_start: (float) temperature at which animation starts
            temperature_stop: (float) temperature at which animation ends
            frames: (int) number of frames in animation
            num_cpu: from calculate_mag_3d
            colour_map_name: name of colour map or list of colours (input for Compound.colour_map)
            lim_scalar: (float) number by which limit of x, y and z axes will be scaled, 1 is default
            ticks: (int) spacing of major ticks on x, y and z axes
            axis_off: (bool) if True axes are not shown
            r_density: (int) density of vertical lines building 3d image
            c_density:(int) density of horizontal lines building 3d image
            filename: (str) name of output file
            fps: (int) number of frames per second in animation
            dpi: (int) resolution of image (dots per inch)

        Returns:

        """

        temps = np.linspace(temperature_start, temperature_stop, frames)
        fig = plt.figure()
        ax = fig.add_subplot(projection='3d')

        writer = PillowWriter(fps=fps)
        xo, yo, zo = self.calculate_mag_3d(group, states_cutoff, field, spherical_grid,
                                           temps, num_cpu)
        with writer.saving(fig, f'{filename}.gif', dpi):
            for temp in range(temps.shape[0]):
                x, y, z = xo[temp], yo[temp], zo[temp]
                max_array = np.array([np.max(x), np.max(y), np.max(z)])
                lim = np.max(max_array)
                norm = plt.Normalize(z.min(), z.max())
                colors = Compound.colour_map(colour_map_name)(norm(z))
                rcount, ccount, _ = colors.shape
                if not r_density:
                    r_density = rcount
                if not c_density:
                    c_density = ccount
                surface = ax.plot_surface(x, y, z, rcount=r_density, ccount=c_density, facecolors=colors, shade=False, )
                surface.set_facecolor((0, 0, 0, 0))
                ax.set_xlim(-lim * lim_scalar, lim * lim_scalar)
                ax.set_ylim(-lim * lim_scalar, lim * lim_scalar)
                ax.set_zlim(-lim * lim_scalar, lim * lim_scalar)
                if ticks == 0:
                    for axis in [ax.xaxis, ax.yaxis, ax.zaxis]:
                        axis.set_ticklabels([])
                        axis._axinfo['axisline']['linewidth'] = 1
                        axis._axinfo['axisline']['color'] = (0, 0, 0)
                        axis._axinfo['grid']['linewidth'] = 0.5
                        axis._axinfo['grid']['linestyle'] = "-"
                        axis._axinfo['grid']['color'] = (0, 0, 0)
                        axis._axinfo['tick']['inward_factor'] = 0.0
                        axis._axinfo['tick']['outward_factor'] = 0.0
                        axis.set_pane_color((0.95, 0.95, 0.95))
                else:
                    ax.xaxis.set_minor_locator(AutoMinorLocator(2))
                    ax.yaxis.set_minor_locator(AutoMinorLocator(2))
                    ax.zaxis.set_minor_locator(AutoMinorLocator(2))
                    ax.xaxis.set_major_locator(MultipleLocator(ticks))
                    ax.yaxis.set_major_locator(MultipleLocator(ticks))
                    ax.zaxis.set_major_locator(MultipleLocator(ticks))
                ax.grid(False)
                ax.set_xlabel(r'$M\ /\ \mathrm{\mu_{B}}$', labelpad=10 * len(str(ticks)) / 4)
                ax.set_ylabel(r'$M\ /\ \mathrm{\mu_{B}}$', labelpad=10 * len(str(ticks)) / 4)
                ax.set_zlabel(r'$M\ /\ \mathrm{\mu_{B}}$', labelpad=10 * len(str(ticks)) / 4)
                ax.set_box_aspect([1, 1, 1])
                if axis_off:
                    plt.axis('off')
                writer.grab_frame()
                plt.cla()

    def axes_in_mol2(self, group, mol2_file, atom_name, scale_factor=1):
        axes_matrix = self[f'{group}_g_tensors_axes', f'{group}_axes'][0]
        g_tensor = self[f'{group}_g_tensors_axes', f'{group}_g_tensors'][0]
        x = axes_matrix[:, 0].T * g_tensor[1]
        y = axes_matrix[:, 1].T * g_tensor[2]
        z = axes_matrix[:, 2].T * g_tensor[3]

        with open(f'{mol2_file}.mol2', 'r', encoding='UTF-8') as mol2:
            file_contents = mol2.readlines()
        reused_information = {'atoms_start': False, 'atoms_end': False}
        for index, line in enumerate(file_contents):

            if '@<TRIPOS>MOLECULE' in line:
                bonds_atoms_count = file_contents[index + 2].split()
                reused_information['number_atoms'] = int(bonds_atoms_count[0])
                reused_information['number_bonds'] = int(bonds_atoms_count[1])
                reused_information['number_molecules'] = int(bonds_atoms_count[2])
                bonds_atoms_count[0] = str(int(bonds_atoms_count[0]) + 6)  # number of atoms added
                bonds_atoms_count[1] = str(int(bonds_atoms_count[1]) + 3)  # number of bonds added
                bonds_atoms_count[2] = str(int(bonds_atoms_count[2]) + 3)  # number of molecules added
                new_contents = ''
                for inner_index in range(len(bonds_atoms_count)):
                    new_contents += '    ' + f'{bonds_atoms_count[inner_index]}'
                file_contents[index + 2] = new_contents
            if '@<TRIPOS>ATOM' in line:
                reused_information['atoms_start'] = True
                continue
            if '@<TRIPOS>BOND' in line:
                reused_information['atoms_end'] = True
            if reused_information['atoms_start']:
                if reused_information['atoms_end'] == False:
                    atom_information = line.split()
                    if atom_information[1] == atom_name:
                        reused_information['atom_x'] = float(atom_information[2])
                        reused_information['atom_y'] = float(atom_information[3])
                        reused_information['atom_z'] = float(atom_information[4])
                        coordinates = np.array(
                            [reused_information['atom_x'], reused_information['atom_y'], reused_information['atom_z']])
                    if int(atom_information[0]) == reused_information['number_atoms']:
                        Xe1 = coordinates + x * scale_factor
                        Xe2 = coordinates - x * scale_factor
                        Fr1 = coordinates + y * scale_factor
                        Fr2 = coordinates - y * scale_factor
                        Rn1 = coordinates + z * scale_factor
                        Rn2 = coordinates - z * scale_factor
                        file_contents[
                            index] += f'    {reused_information["number_atoms"] + 1} Xe1     {Xe1[0]}   {Xe1[1]}   {Xe1[2]}   Xe        {reused_information["number_molecules"] + 1} RES{reused_information["number_molecules"] + 1}   0.0000\n    {reused_information["number_atoms"] + 2} Xe2     {Xe2[0]}   {Xe2[1]}   {Xe2[2]}   Xe        {reused_information["number_molecules"] + 1} RES{reused_information["number_molecules"] + 1}   0.0000\n    {reused_information["number_atoms"] + 3} Fr1     {Fr1[0]}   {Fr1[1]}   {Fr1[2]}   Fr        {reused_information["number_molecules"] + 2} RES{reused_information["number_molecules"] + 2}   0.0000\n    {reused_information["number_atoms"] + 4} Fr2     {Fr2[0]}   {Fr2[1]}   {Fr2[2]}   Fr        {reused_information["number_molecules"] + 2} RES{reused_information["number_molecules"] + 2}   0.0000\n    {reused_information["number_atoms"] + 5} Rn1     {Rn1[0]}   {Rn1[1]}   {Rn1[2]}   Rn        {reused_information["number_molecules"] + 3} RES{reused_information["number_molecules"] + 3}   0.0000\n    {reused_information["number_atoms"] + 6} Rn2     {Rn2[0]}   {Rn2[1]}   {Rn2[2]}   Rn        {reused_information["number_molecules"] + 3} RES{reused_information["number_molecules"] + 3}   0.0000\n'

            if '@<TRIPOS>SUBSTRUCTURE' in line:
                file_contents[
                    index] = f'    {reused_information["number_bonds"] + 1}    {reused_information["number_atoms"] + 1}    {reused_information["number_atoms"] + 2}   1\n    {reused_information["number_bonds"] + 2}    {reused_information["number_atoms"] + 3}    {reused_information["number_atoms"] + 4}   1\n    {reused_information["number_bonds"] + 3}    {reused_information["number_atoms"] + 5}    {reused_information["number_atoms"] + 6}   1\n@<TRIPOS>SUBSTRUCTURE\n'
            if '@<TRIPOS>CRYSIN' in line:
                file_contents[
                    index] = f'     {reused_information["number_molecules"] + 1} RES{reused_information["number_molecules"] + 1}       {reused_information["number_atoms"] + 1} GROUP             0 ****  ****    0 \n     {reused_information["number_molecules"] + 2} RES{reused_information["number_molecules"] + 2}       {reused_information["number_atoms"] + 3} GROUP             0 ****  ****    0 \n     {reused_information["number_molecules"] + 3} RES{reused_information["number_molecules"] + 3}       {reused_information["number_atoms"] + 5} GROUP             0 ****  ****    0 \n@<TRIPOS>CRYSIN\n'
        for line in file_contents:
            with open(f'{mol2_file}_axes.mol2', 'a', encoding='UTF-8') as mol2:
                mol2.write(line)
