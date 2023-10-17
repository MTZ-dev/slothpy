"""Module for experimental exports of SlothPy data."""
from typing import Union
from os.path import join

from pandas import DataFrame, read_csv, concat

from docx import Document as Doc
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from numpy import array, float64

from slothpy.core.compound_object import Compound
from slothpy.core._slothpy_exceptions import (
    SltFileError,
    SltSaveError,
    SltInputError,
)
from slothpy._general_utilities._constants import BLUE, YELLOW, RESET


def table_energy_and_g(
    slt_file: Compound,
    group: str,
    i_last_doublet: int,
    krames: bool,
    i_first_doublet: int = 0,
    decomp: Union["total_angular", "magnetic"] = "total_angular",
    i_first_composition: int = 0,
    i_last_composition: int = 0,
    threshold: float = 0.1,
    output_path: str = "./",
    output_name: str = "",
):
    """
    Creates .docx file containing table with g tensor and energy information.

    Parameters
    ----------
    slt_file : Compound
        Slt file storing data needed for the table.
    group: str
        Name of a group from .slt file for which table will be created.
        Requires f"{group}_soc_energies", f"{group}_g_tensors_axes" and
        f"{group}_magnetic_decomposition" or f"{group}_total_angular
        _decomposition" (outputs of Compound functions: soc_energies_cm_1,
        calculate_g_tensor_and_axes_doublet, matrix_decomposition_in_z_pseudo
        _spin_basis(with 'soc' and 'total_angular' or 'magnetic' setting,
        additionally rotation should be set as one for the ground state)).
    i_last_doublet: int
        Index of the last doublet used in the table. For example, if you are
        working with the lanthanides ions 3+ this index should be matching the
        splitting of the ground term of the considered lanthanide.
        (J2+1 states, since we take index of doublet /2 and -1 since we count
        from 0): Ce: 2, Pr: 3, Nd: 4, Pm: 3, Sm: 2, Eu: 0, Gd: 3, Yb: 3, Tm: 5,
        Er: 7, Ho: 7, Dy: 7, Tb: 5
    krames: bool
        Determines if the table should use the table's template for Krames
        ions (even electron number) or non-Krames ions (odd electron number).
    i_first_doublet: int = 0
        Index of the last doublet used in the table.
    decomp: Union["total_angular", "magnetic"] = "total_angular"
        Determines the type of decomposition used.
    i_first_composition: int = 0
        Determines first decomposed doublet, if 0 i_first_composition = i_
        first_doublet.
    i_last_composition: int = 0
        Determines last decomposed level, if 0
        i_last_composition = i_last_doublet.
    threshold: float = 0.1
        Determines how high the contribution of energy state has to be to be
        shown in the table [%].
    output_path: str = "./"
        Path to which the output will be saved.
    output_name: str = ""
        Name of the output .docx file.

    Returns
    -------
    Nothing

    Raises
    ------
    SltFileError
        If a Compound object is not passed as an input or it doesn't include
        all the necessary data.
    SltSaveError
        If the program is unable to correctly save the .docx file.

    See Also
    --------
    slothpy.Compound.calculate_g_tensor_and_axes_doublet
    slothpy.Compound.soc_energies_cm_1
    slothpy.Compound.matrix_decomposition_in_z_pseudo_spin_basis

    """
    if not isinstance(slt_file, Compound):
        raise SltFileError(
            TypeError(""), messege="A Compound object must be passed!"
        ) from None

    if not i_first_composition:
        i_first_composition = i_first_doublet
    if not i_last_composition:
        i_last_composition = i_last_doublet
    if not output_name:
        output_name = f"{group}_table_energy_and_g"

    if not all(
        isinstance(var, int)
        for var in (
            i_last_doublet,
            i_first_doublet,
            i_first_composition,
            i_last_composition,
        )
    ):
        raise SltInputError(
            ValueError("All passed indexes have to be integers.")
        )

    try:
        energies = slt_file[f"{group}_soc_energies", f"{group}_energies"]
        g_tensors = slt_file[f"{group}_g_tensors_axes", f"{group}_g_tensors"]
        if decomp == "total_angular":
            composition_frac_matrix = slt_file[
                f"{group}_total_angular_decomposition",
                f"{group}_total_angular_decomposition",
            ]
            composition_states = slt_file[
                f"{group}_total_angular_decomposition",
                f"{group}_pseudo_spin_states",
            ]
        elif decomp == "magnetic":
            composition_frac_matrix = slt_file[
                f"{group}_magnetic_decomposition",
                f"{group}_magnetic_decomposition",
            ]
            composition_states = slt_file[
                f"{group}_magnetic_decomposition",
                f"{group}_pseudo_spin_states",
            ]
    except Exception as exc:
        raise SltFileError(
            slt_file._hdf5,
            exc,
            f"Failed to load data required by table"
            + BLUE
            + "Group "
            + RESET
            + '"'
            + BLUE
            + f"{group}"
            + RESET
            + '".'
            + RED
            + "Check if group exist.",
        ) from None

    try:
        number_of_doublets = i_last_doublet - i_first_doublet + 1
        number_to_decompose = i_last_composition - i_first_composition + 1
        docx = Doc()
        if krames:
            table = docx.add_table(rows=4 + number_of_doublets, cols=5)
            row = 0

            table.cell(row, 0).paragraphs[0].add_run(f"{group}").bold = True
            table.cell(row, 0).merge(table.cell(row, 4))
            row += 1

            table.cell(row, 0).paragraphs[0].add_run(f"Energy and pseudo-")
            table.cell(row, 0).paragraphs[0].add_run(f"g").italic = True
            table.cell(row, 0).paragraphs[0].add_run("-tensor components (")
            table.cell(row, 0).paragraphs[0].add_run(f"g").italic = True
            subscript = table.cell(row, 0).paragraphs[0].add_run("x")
            subscript.font.subscript = True
            table.cell(row, 0).paragraphs[0].add_run(", ")
            table.cell(row, 0).paragraphs[0].add_run(f"g").italic = True
            subscript = table.cell(row, 0).paragraphs[0].add_run("y")
            subscript.font.subscript = True
            table.cell(row, 0).paragraphs[0].add_run(", ")
            table.cell(row, 0).paragraphs[0].add_run(f"g").italic = True
            subscript = table.cell(row, 0).paragraphs[0].add_run("z")
            subscript.font.subscript = True
            table.cell(row, 0).paragraphs[0].add_run(
                ") of"
                f" {number_of_doublets}{' ground' if i_first_doublet == 0 else ''} Kramers"
                " doublets"
            )
            table.cell(row, 0).merge(table.cell(row, 4))
            row += 1

            table.cell(row, 0).paragraphs[0].add_run(f"Doublet no.")

            table.cell(row, 1).paragraphs[0].add_run(f"Energy / cm")
            superscript = table.cell(row, 1).paragraphs[0].add_run("-1")
            superscript.font.superscript = True

            table.cell(row, 2).paragraphs[0].add_run(f"Pseudo-")
            table.cell(row, 2).paragraphs[0].add_run("g").italic = True
            table.cell(row, 2).paragraphs[0].add_run("-tensor components")
            table.cell(row, 2).merge(table.cell(row, 4))
            row += 1

            table.cell(row - 1, 0).merge(table.cell(row, 0))
            table.cell(row - 1, 1).merge(table.cell(row, 1))

            table.cell(row, 2).paragraphs[0].add_run("g").italic = True
            subscript = table.cell(row, 2).paragraphs[0].add_run("x")
            subscript.font.subscript = True

            table.cell(row, 3).paragraphs[0].add_run("g").italic = True
            subscript = table.cell(row, 3).paragraphs[0].add_run("y")
            subscript.font.subscript = True

            table.cell(row, 4).paragraphs[0].add_run("g").italic = True
            subscript = table.cell(row, 4).paragraphs[0].add_run("z")
            subscript.font.subscript = True
            row += 1

            for index in range(number_of_doublets):
                table.cell(row, 0).paragraphs[0].add_run(
                    f"{int(g_tensors[i_first_doublet + index][0]) + 1}."
                )
                table.cell(row, 1).paragraphs[0].add_run(
                    f"{energies[(i_first_doublet + index) * 2]:.3f}"
                )
                table.cell(row, 2).paragraphs[0].add_run(
                    f"{g_tensors[i_first_doublet + index][1]:.4f}"
                )
                table.cell(row, 3).paragraphs[0].add_run(
                    f"{g_tensors[i_first_doublet + index][2]:.4f}"
                )
                table.cell(row, 4).paragraphs[0].add_run(
                    f"{g_tensors[i_first_doublet + index][3]:.4f}"
                )
                row += 1

            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row1 in table.rows:
                for cell in row1.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        elif not krames:
            table = docx.add_table(rows=3 + number_of_doublets, cols=4)
            row = 0

            table.cell(row, 0).paragraphs[0].add_run(f"{group}").bold = True
            table.cell(row, 0).merge(table.cell(row, 3))
            row += 1

            table.cell(row, 0).paragraphs[0].add_run(f"Energy and pseudo-")
            table.cell(row, 0).paragraphs[0].add_run(f"g").italic = True
            table.cell(row, 0).paragraphs[0].add_run("-tensor")
            table.cell(row, 0).paragraphs[0].add_run(f" g").italic = True
            subscript = table.cell(row, 0).paragraphs[0].add_run("z")
            subscript.font.subscript = True
            table.cell(row, 0).paragraphs[0].add_run(f" component")
            table.cell(row, 0).paragraphs[0].add_run(
                f" of {number_of_doublets}{' ground' if i_first_doublet == 0 else ''} Ising"
                " doublets"
            )
            table.cell(row, 0).merge(table.cell(row, 3))
            row += 1

            table.cell(row, 0).paragraphs[0].add_run(f"Doublet no.")

            table.cell(row, 1).paragraphs[0].add_run(f"Energies / cm")
            superscript = table.cell(row, 1).paragraphs[0].add_run("-1")
            superscript.font.superscript = True

            table.cell(row, 2).paragraphs[0].add_run(
                f"Tunneling splitting / cm"
            )
            superscript = table.cell(row, 2).paragraphs[0].add_run("-1")
            superscript.font.superscript = True

            table.cell(row, 3).paragraphs[0].add_run(f"g").italic = True
            subscript = table.cell(row, 3).paragraphs[0].add_run("z")
            subscript.font.subscript = True
            table.cell(row, 3).paragraphs[0].add_run(f" component")
            row += 1

            for index in range(number_of_doublets):
                table.cell(row, 0).paragraphs[0].add_run(
                    f"{int(g_tensors[i_first_doublet + index][0]) + 1}."
                )
                table.cell(row, 1).paragraphs[0].add_run(
                    f"{energies[(i_first_doublet + index * 2)]:.3f}; "
                )
                table.cell(row, 1).paragraphs[0].add_run(
                    f"{energies[(i_first_doublet + 1 + index * 2)]:.3f}"
                )
                table.cell(row, 2).paragraphs[0].add_run(
                    f"{abs(energies[(i_first_doublet + 1 + index * 2)] - energies[(i_first_doublet + index * 2)]):.3e}"
                )
                table.cell(row, 3).paragraphs[0].add_run(
                    f"{g_tensors[i_first_doublet + index][2]:.4f}"
                )
                row += 1

            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for row1 in table.rows:
                for cell in row1.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        table2 = docx.add_table(rows=3, cols=number_to_decompose)
        table2.cell(0, 0).paragraphs[0].add_run(
            "Composition of the two ground"
            f" {'Kramers' if krames else 'Ising'} doublets in the |"
        )
        table2.cell(0, 0).paragraphs[0].add_run("m").italic = True
        subscript = table2.cell(0, 0).paragraphs[0].add_run("J")
        subscript.font.subscript = True
        table2.cell(0, 0).paragraphs[0].add_run(
            "⟩ basis on the quantization axes within"
        )
        table2.cell(0, 0).paragraphs[0].add_run(" J ").italic = True
        if krames:
            table2.cell(0, 0).paragraphs[0].add_run(
                f"= {str(int(abs(composition_states[0] * 2)))}/2 manifold"
            )
        else:
            table2.cell(0, 0).paragraphs[0].add_run(
                f"= {str(int(abs(composition_states[0])))} manifold"
            )
        table2.cell(0, 0).paragraphs[0].add_run(
            f" (contribution over {threshold}% shown)"
        )
        table2.cell(0, 0).merge(table2.cell(0, number_to_decompose - 1))

        for index in range(number_to_decompose):
            table2.cell(1, index).paragraphs[0].add_run(f"{index + 1}")
            if index + 1 == 1:
                superscript = table2.cell(1, index).paragraphs[0].add_run("st")
                superscript.font.superscript = True
            elif index + 1 == 2:
                superscript = table2.cell(1, index).paragraphs[0].add_run("nd")
                superscript.font.superscript = True
            elif index + 1 == 3:
                superscript = table2.cell(1, index).paragraphs[0].add_run("rd")
                superscript.font.superscript = True
            else:
                superscript = table2.cell(1, index).paragraphs[0].add_run("th")
                superscript.font.superscript = True
            table2.cell(1, index).paragraphs[0].add_run(" doublet")
            for inner_index in range(len(composition_states)):
                if (
                    composition_frac_matrix[i_first_composition + index * 2][
                        inner_index
                    ]
                    <= threshold
                ):
                    continue
                else:
                    table2.cell(2, index).paragraphs[0].add_run(
                        f"{composition_frac_matrix[(i_first_composition + index * 2)][inner_index]:.1f}%"
                    )
                    table2.cell(2, index).paragraphs[0].add_run(
                        f' |{"+" + str(int(abs(composition_states[inner_index] * 2))) if composition_states[inner_index] > 0 else "–" + str(int(abs(composition_states[inner_index] * 2)))}/2⟩\n'
                    )
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row2 in table2.rows:
            for cell in row2.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        docx.save(join(output_path, f"{output_name}.docx"))

        print(
            YELLOW
            + "SlothPy message: "
            + RESET
            + f"Table was saved as {output_name} in {output_path} as docx"
            " file."
        )

    except Exception as exc:
        raise SltSaveError(
            slt_file._hdf5,
            exc,
            f"Failed to save table {output_name} from "
            + BLUE
            + "Group "
            + RESET
            + '"'
            + BLUE
            + f"{group}"
            + RESET
            + f'"in {output_path}.',
        ) from None


def axes_in_xyz(
    slt_file: Compound,
    group: str,
    central_atom: str,
    xyz_path: str,
    xyz_file_name: str,
    doublet_number: int = 0,
    scale_factor: float64 = 1,
    output_path: str = "./",
):
    """
    Adds main XYZ magnetic axes corresponding to the chosen doublet
    to a .xyz file.

    Parameters
    ----------
    slt_file : Compound
        Name of a Compound object corresponding to the .slt file that will
        be used.
    group : str
        Name of the group with axes and g_tensors (result of the
        calculate_g_tensor_and_axes_doublet method).
    xyz_path : str
        Path to the .xyz file to which axes will be added.
    xyz_file_name : str
        Name of the .xyz file to which axes will be added.
    central_atom : str
        Symbol of a unique central atom from which axes will begin.
    doublet_number : int, optional
        Number of the doublet whose magnetic axes will be added., by default 0
    scale_factor : float64, optional
        Factor by which the lengths of the axes will be scaled., by default 1
    output_path : str, optional
        Path to which .xyz output file with the suffix _axes will be saved.,
        by default "./"

    Raises
    ------
    SltFileError
        If a Compound object is not passed as an input.
    SltInputError
        If the name of an input xyz file does not end with .xyz.
    SltSaveError
        If the program is unable to correctly save the new .xyz file.

    Note
    ----
    Magnetic axes are scaled by the corresponding pseudo-g-tensor values along
    them. Atoms representing the axes are as follows: Lv - X, Ts - Y, Og - Z.

    See Also
    --------
    slothpy.Compound.calculate_g_tensor_and_axes_doublet
    """
    if not isinstance(slt_file, Compound):
        raise SltFileError(
            TypeError(""), messege="A Compound object must be passed!"
        ) from None

    if xyz_file_name[-4:] != ".xyz":
        raise SltInputError(NameError("Input file name has to end with .xyz."))

    try:
        axes_matrix = slt_file[f"{group}_g_tensors_axes", f"{group}_axes"][
            doublet_number
        ]
        g_tensor = slt_file[f"{group}_g_tensors_axes", f"{group}_g_tensors"][
            doublet_number
        ]

    except Exception as exc:
        raise SltFileError(
            slt_file._hdf5,
            exc,
            f"Failed to load data required to add axes"
            + BLUE
            + "Group "
            + RESET
            + '"'
            + BLUE
            + f"{group}"
            + RESET
            + '".'
            + RED
            + "Check if group exist.",
        ) from None

    try:
        x = axes_matrix[:, 0].T * g_tensor[1] * scale_factor
        y = axes_matrix[:, 1].T * g_tensor[2] * scale_factor
        z = axes_matrix[:, 2].T * g_tensor[3] * scale_factor

        atom_dict = {"Lv": x, "Ts": y, "Og": z}

        # Read the XYZ file into a DataFrame
        xyz_file = join(xyz_path, xyz_file_name)
        atoms_df = read_csv(
            xyz_file,
            delim_whitespace=True,
            skiprows=2,
            header=None,
            names=["Element", "X", "Y", "Z"],
        )
        central_atom_coord = atoms_df[atoms_df["Element"] == central_atom][
            ["X", "Y", "Z"]
        ]
        new_atom_names = ["Lv", "Ts", "Og"]

        # Create a DataFrame with the new atoms with Element names
        for new_atom in new_atom_names:
            new_atom_plus = DataFrame({"Element": [new_atom]})
            new_atom_plus[["X", "Y", "Z"]] = (
                central_atom_coord + atom_dict[new_atom]
            )
            atoms_df = concat([atoms_df, new_atom_plus], ignore_index=True)

            new_atom_minus = DataFrame({"Element": [new_atom]})
            new_atom_minus[["X", "Y", "Z"]] = (
                central_atom_coord - atom_dict[new_atom]
            )
            atoms_df = concat([atoms_df, new_atom_minus], ignore_index=True)

        output_name = f"{xyz_file_name[:-4]}_axes.xyz"
        # Save the updated XYZ file
        output_xyz_file = join(output_path, output_name)
        with open(output_xyz_file, "w") as f:
            f.write(f"{len(atoms_df)}\n")
            f.write(
                "Generated by SlothPy - magnetic axes (Lv - X, Ts - Y, Og - Z)"
                " scaled by the corresponding pseudo-g-tensors.\n"
            )
            atoms_df.to_csv(
                f,
                sep="\t",
                header=False,
                index=False,
                float_format="%.7f",
                encoding="utf-8",
            )

        print(
            YELLOW
            + "SlothPy message: "
            + RESET
            + f"Updated .xyz file was saved as {output_name} in {output_path}."
        )

    except Exception as exc:
        raise SltSaveError(
            slt_file._hdf5,
            exc,
            f"Failed to save new .xyz file {output_name} with magnetic axes"
            " from "
            + BLUE
            + "Group "
            + RESET
            + '"'
            + BLUE
            + f"{group}"
            + RESET
            + f'"in {output_path}.',
        ) from None


def axes_in_mol2(
    slt_file: Compound,
    group: str,
    mol2_file_path: str,
    mol2_file_name: str,
    atom_name: str,
    doublet_number: int = 0,
    scale_factor: float = 1,
    output_path: str = "",
):
    """
    Adds main XYZ magnetic axes corresponding to the chosen doublet
    to a .mol2 file.

    Parameters
    ----------
    slt_file : Compound
        Name of a Compound object corresponding to the .slt file that will
        be used.
    group : str
        Name of the group with axes and g_tensors (result of the
        calculate_g_tensor_and_axes_doublet method).
    mol2_file_path : str
        Path to the .mol2 file to which axes will be added.
    mol2_file_name : str
        Name of the .mol2 file to which axes will be added.
    atom_name : str
        Name of a central atom from which axes will begin.
    doublet_number : int = 0
        Number of the doublet whose magnetic axes will be added., by default 0
    scale_factor : float64 = 1
        Factor by which the lengths of the axes will be scaled., by default 1
    output_path : str = "./"
        Path to which .mol2 output file with the suffix _axes will be saved.,
        by default "./"

    Raises
    ------
    SltFileError
        If a Compound object is not passed as an input.
    SltInputError
        If the name of an input xyz file does not end with .xyz.
    SltSaveError
        If the program is unable to correctly save the new .xyz file.

    Note
    ----
    Magnetic axes are scaled by the corresponding pseudo-g-tensor values along
    them. Atoms representing the axes are as follows: Lv - X, Ts - Y, Og - Z.

    See Also
    --------
    slothpy.Compound.calculate_g_tensor_and_axes_doublet
    """

    if not isinstance(slt_file, Compound):
        raise SltFileError(
            TypeError(""), messege="A Compound object must be passed!"
        ) from None

    if mol2_file_name[-5:] != ".mol2":
        raise SltInputError(
            NameError("Input file name has to end with .mol2.")
        )

    if not output_path:
        output_path = mol2_file_path
    output_name = f"{mol2_file_name[:-5]}_axes.mol2"

    try:
        axes_matrix = slt_file[f"{group}_g_tensors_axes", f"{group}_axes"][
            doublet_number
        ]
        g_tensor = slt_file[f"{group}_g_tensors_axes", f"{group}_g_tensors"][
            doublet_number
        ]
    except Exception as exc:
        raise SltFileError(
            slt_file._hdf5,
            exc,
            f"Failed to load data required to add axes"
            + BLUE
            + "Group "
            + RESET
            + '"'
            + BLUE
            + f"{group}"
            + RESET
            + '".'
            + RED
            + "Check if group exist.",
        ) from None

    x = axes_matrix[:, 0].T * g_tensor[1]
    y = axes_matrix[:, 1].T * g_tensor[2]
    z = axes_matrix[:, 2].T * g_tensor[3]

    try:
        with open(
            join(mol2_file_path, mol2_file_name), "r", encoding="UTF-8"
        ) as mol2:
            file_contents = mol2.readlines()
        reused_information = {"atoms_start": False, "atoms_end": False}
        for index, line in enumerate(file_contents):
            if "@<TRIPOS>MOLECULE" in line:
                bonds_atoms_count = file_contents[index + 2].split()
                reused_information["number_atoms"] = int(bonds_atoms_count[0])
                reused_information["number_bonds"] = int(bonds_atoms_count[1])
                reused_information["number_molecules"] = int(
                    bonds_atoms_count[2]
                )
                bonds_atoms_count[0] = str(
                    int(bonds_atoms_count[0]) + 6
                )  # number of atoms added
                bonds_atoms_count[1] = str(
                    int(bonds_atoms_count[1]) + 3
                )  # number of bonds added
                bonds_atoms_count[2] = str(
                    int(bonds_atoms_count[2]) + 3
                )  # number of molecules added
                new_contents = ""
                for inner_index in range(len(bonds_atoms_count)):
                    new_contents += (
                        "    " + f"{bonds_atoms_count[inner_index]}"
                    )
                file_contents[index + 2] = new_contents
            if "@<TRIPOS>ATOM" in line:
                reused_information["atoms_start"] = True
                continue
            if "@<TRIPOS>BOND" in line:
                reused_information["atoms_end"] = True
            if reused_information["atoms_start"]:
                if reused_information["atoms_end"] == False:
                    atom_information = line.split()
                    if atom_information[1] == atom_name:
                        reused_information["atom_x"] = float(
                            atom_information[2]
                        )
                        reused_information["atom_y"] = float(
                            atom_information[3]
                        )
                        reused_information["atom_z"] = float(
                            atom_information[4]
                        )
                        coordinates = array(
                            [
                                reused_information["atom_x"],
                                reused_information["atom_y"],
                                reused_information["atom_z"],
                            ]
                        )
                    if (
                        int(atom_information[0])
                        == reused_information["number_atoms"]
                    ):
                        Lv1 = coordinates + x * scale_factor
                        Lv2 = coordinates - x * scale_factor
                        Ts1 = coordinates + y * scale_factor
                        Ts2 = coordinates - y * scale_factor
                        Og1 = coordinates + z * scale_factor
                        Og2 = coordinates - z * scale_factor
                        file_contents[index] += (
                            f'    {reused_information["number_atoms"] + 1} Lv1'
                            f"     {Lv1[0]}   {Lv1[1]}   {Lv1[2]}   Lv       "
                            f' {reused_information["number_molecules"] + 1} RES{reused_information["number_molecules"] + 1} '
                            "  0.0000\n   "
                            f' {reused_information["number_atoms"] + 2} Lv2   '
                            f"  {Lv2[0]}   {Lv2[1]}   {Lv2[2]}   Lv       "
                            f' {reused_information["number_molecules"] + 1} RES{reused_information["number_molecules"] + 1} '
                            "  0.0000\n   "
                            f' {reused_information["number_atoms"] + 3} Ts1   '
                            f"  {Ts1[0]}   {Ts1[1]}   {Ts1[2]}   Ts       "
                            f' {reused_information["number_molecules"] + 2} RES{reused_information["number_molecules"] + 2} '
                            "  0.0000\n   "
                            f' {reused_information["number_atoms"] + 4} Ts2   '
                            f"  {Ts2[0]}   {Ts2[1]}   {Ts2[2]}   Ts       "
                            f' {reused_information["number_molecules"] + 2} RES{reused_information["number_molecules"] + 2} '
                            "  0.0000\n   "
                            f' {reused_information["number_atoms"] + 5} Og1   '
                            f"  {Og1[0]}   {Og1[1]}   {Og1[2]}   Og       "
                            f' {reused_information["number_molecules"] + 3} RES{reused_information["number_molecules"] + 3} '
                            "  0.0000\n   "
                            f' {reused_information["number_atoms"] + 6} Og2   '
                            f"  {Og2[0]}   {Og2[1]}   {Og2[2]}   Og       "
                            f' {reused_information["number_molecules"] + 3} RES{reused_information["number_molecules"] + 3} '
                            "  0.0000\n"
                        )

            if "@<TRIPOS>SUBSTRUCTURE" in line:
                file_contents[index] = (
                    f'    {reused_information["number_bonds"] + 1}   '
                    f' {reused_information["number_atoms"] + 1}   '
                    f' {reused_information["number_atoms"] + 2}   1\n   '
                    f' {reused_information["number_bonds"] + 2}   '
                    f' {reused_information["number_atoms"] + 3}   '
                    f' {reused_information["number_atoms"] + 4}   1\n   '
                    f' {reused_information["number_bonds"] + 3}   '
                    f' {reused_information["number_atoms"] + 5}   '
                    f' {reused_information["number_atoms"] + 6}  '
                    " 1\n@<TRIPOS>SUBSTRUCTURE\n"
                )
            if "@<TRIPOS>CRYSIN" in line:
                file_contents[index] = (
                    "    "
                    f" {reused_information['number_molecules'] + 1} RES{reused_information['number_molecules'] + 1} "
                    f"      {reused_information['number_atoms'] + 1} GROUP    "
                    "         0 ****  ****    0 \n    "
                    f" {reused_information['number_molecules'] + 2} RES{reused_information['number_molecules'] + 2} "
                    f"      {reused_information['number_atoms'] + 3} GROUP    "
                    "         0 ****  ****    0 \n    "
                    f" {reused_information['number_molecules'] + 3} RES{reused_information['number_molecules'] + 3} "
                    f"      {reused_information['number_atoms'] + 5} GROUP    "
                    "         0 ****  ****    0 \n@<TRIPOS>CRYSIN\n"
                )
        for line in file_contents:
            with open(
                join(output_path, output_name), "a", encoding="UTF-8"
            ) as mol2:
                mol2.write(line)
        print(
            YELLOW
            + "SlothPy message: "
            + RESET
            + f"Updated .mol2 file was saved as {output_name} in"
            f" {output_path}."
        )

    except Exception as exc:
        raise SltSaveError(
            slt_file._hdf5,
            exc,
            f"Failed to save new .mol2 file {output_name} with magnetic axes"
            " from "
            + BLUE
            + "Group "
            + RESET
            + '"'
            + BLUE
            + f"{group}"
            + RESET
            + f'"in {output_path}.',
        ) from None
